import json
import shutil
import tempfile
from zipfile import ZipFile, ZIP_DEFLATED
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
TEMPLATE = Path(r"C:\Users\Justin\Desktop\LinkedIn\Target_Product_Profile_cohere_report.docx")
OUTPUT = ROOT / "Vestrum_DME_Methods_Cohere_Report.docx"
DETAILED_JSON = ROOT / "methods_extraction_detailed.json"
SOURCE_LOG_JSON = ROOT / "source_log.json"

INK = "10120F"
CREAM = "F6F1E8"
MUTED = "5C6257"
LIME = "D7FF5F"
BORDER = "10120F"
LINK = "2B5C8A"


def load_data():
    return json.loads(DETAILED_JSON.read_text(encoding="utf-8")), json.loads(
        SOURCE_LOG_JSON.read_text(encoding="utf-8")
    )


def clear_document(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_section(section, landscape=False):
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.45)
        section.right_margin = Inches(0.45)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)


def set_run(run, size=None, bold=None, color=None, italic=None):
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text, url, color=LINK, underline=True, size=None, bold=False):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    if size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))
        r_pr.append(sz)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")


def set_cell_borders(cell, color=BORDER, size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_pr(table, width_dxa="10080", border_color="1B1F17", border_size="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is not None:
        tbl_pr.remove(tbl_layout)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), width_dxa)

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), border_size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), border_color)

    cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for side in ("left", "right"):
        element = cell_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            cell_mar.append(element)
        element.set(qn("w:w"), "10")
        element.set(qn("w:type"), "dxa")

    tbl_look = tbl_pr.find(qn("w:tblLook"))
    if tbl_look is None:
        tbl_look = OxmlElement("w:tblLook")
        tbl_pr.append(tbl_look)
    tbl_look.set(qn("w:val"), "0000")
    tbl_look.set(qn("w:firstRow"), "0")
    tbl_look.set(qn("w:lastRow"), "0")
    tbl_look.set(qn("w:firstColumn"), "0")
    tbl_look.set(qn("w:lastColumn"), "0")
    tbl_look.set(qn("w:noHBand"), "0")
    tbl_look.set(qn("w:noVBand"), "0")


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def set_cell_text(cell, text, color=INK, bold=False, size=7.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.1
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_doc_hyperlink_cell(cell, label, url):
    cell.text = ""
    p = cell.paragraphs[0]
    add_hyperlink(p, label, url, size=7.2, bold=False)
    p.paragraph_format.space_after = Pt(0)


def style_table(table, widths=None, font_size=7.2):
    table.autofit = True
    set_table_pr(table)
    mark_header_row(table.rows[0])
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            shade_cell(cell, INK if r_idx == 0 else CREAM)
            set_cell_borders(cell, color="1B1F17", size="4")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    set_run(run, size=font_size, color=LIME if r_idx == 0 else INK, bold=(r_idx == 0))
            if widths and c_idx < len(widths):
                cell.width = Inches(widths[c_idx])


def add_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run(r, size=22, bold=True, color=INK)
    return p


def add_eyebrow(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=8, bold=True, color=INK)
    return p


def add_heading(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"{number}.  {text}")
    set_run(r, size=15, bold=True, color=INK)
    return p


def add_body(doc, text, citations=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.33
    r = p.add_run(text)
    set_run(r, size=10, color=MUTED)
    if citations:
        p.add_run(" ")
        for idx, (label, url) in enumerate(citations):
            if idx:
                p.add_run(" ")
            add_hyperlink(p, f"[{label}]", url, size=9)
    return p


def add_bullet(doc, text, citations=None):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.32)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.18
    tabs = p._p.get_or_add_pPr().find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        p._p.get_or_add_pPr().append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), "461")
    tabs.append(tab)
    r = p.add_run("•\t")
    set_run(r, size=10, color=INK, bold=True)
    r2 = p.add_run(text)
    set_run(r2, size=9.5, color=MUTED)
    if citations:
        p.add_run(" ")
        for idx, (label, url) in enumerate(citations):
            if idx:
                p.add_run(" ")
            add_hyperlink(p, f"[{label}]", url, size=8.5)


def add_callout(doc, title, text, citations=None):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    set_table_pr(table)
    cell = table.cell(0, 0)
    shade_cell(cell, INK)
    set_cell_borders(cell, color="1B1F17", size="4")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run(r, size=8, bold=True, color=LIME)
    p.paragraph_format.space_after = Pt(7)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.33
    r2 = p2.add_run(text)
    set_run(r2, size=9.5, color=CREAM)
    if citations:
        p2.add_run(" ")
        for label, url in citations:
            add_hyperlink(p2, f"[{label}]", url, color=LIME, size=8.5)
    doc.add_paragraph()


def source_url(row):
    return row["source_urls"].split(";")[0].strip()


def table_caption(doc, text, citations=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run(r, size=8.5, bold=True, color=INK)
    if citations:
        p.add_run(" ")
        for label, url in citations:
            add_hyperlink(p, f"[{label}]", url, size=8)


def add_inventory_table(doc, rows):
    table_caption(
        doc,
        "Publication inventory and access basis.",
        [("Vestrum media", "https://www.vestrumhealth.com/media.php"), ("source log", SOURCE_LOG_JSON.as_uri())],
    )
    table = doc.add_table(rows=1, cols=6)
    headers = ["ID", "Scope", "Publication", "Evidence access", "Best source", "Caveat"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, color=LIME, bold=True, size=7.5)
    for row in rows:
        cells = table.add_row().cells
        vals = [
            row["id"],
            "Core" if row["source_scope"].startswith("Core") else "Context",
            f"{row['title']} ({row['year']})",
            row["full_text_status"],
            "",
            row["limitations_relevant_to_methods"],
        ]
        for i, val in enumerate(vals):
            if i == 4:
                add_doc_hyperlink_cell(cells[i], "open source", source_url(row))
            else:
                set_cell_text(cells[i], val, size=6.7)
    style_table(table, widths=[0.55, 0.65, 2.6, 2.1, 0.85, 2.2], font_size=6.7)


def add_methods_table(doc, rows):
    table_caption(
        doc,
        "Operational definitions extracted from full text where available.",
        [("detailed extraction", DETAILED_JSON.as_uri())],
    )
    table = doc.add_table(rows=1, cols=7)
    headers = [
        "ID",
        "DME definition / codes",
        "Treatment definition / codes",
        "Index + follow-up",
        "VA / outcome algorithm",
        "Statistics",
        "Appendix / supplement",
    ]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, color=LIME, bold=True, size=7.2)
    for row in rows:
        cells = table.add_row().cells
        vals = [
            row["id"],
            f"{row['dme_definition_detail']} Codes: {row['diagnostic_codes_reported']}",
            f"{row['treatment_definition_detail']} Codes/source: {row['procedure_drug_codes_reported_or_location']}",
            f"{row['index_date_or_time_origin']} Follow-up: {row['follow_up_and_attrition_rules']}",
            f"{row['va_method']} Outcomes: {row['outcome_operational_definitions']}",
            row["statistical_methods"],
            row["appendix_or_supplement_location"],
        ]
        for i, val in enumerate(vals):
            set_cell_text(cells[i], val, size=5.9)
    style_table(table, widths=[0.45, 1.7, 1.85, 1.7, 2.0, 1.25, 1.55], font_size=5.9)


def add_recommendation_table(doc):
    table_caption(
        doc,
        "Recommended future Vestrum DME study design, based on recurring methods.",
        [
            ("BJO 28,658 eyes", "https://pmc.ncbi.nlm.nih.gov/articles/PMC7848066/"),
            ("Pitcher", "https://pmc.ncbi.nlm.nih.gov/articles/PMC9979048/"),
            ("Sodhi", "https://pmc.ncbi.nlm.nih.gov/articles/PMC10170625/"),
        ],
    )
    rows = [
        [
            "Cohort anchor",
            "Use first anti-VEGF injection for anti-VEGF outcome studies; use new DME diagnosis for treatment-pattern studies.",
            "Aligns with core-04/core-02 for outcomes and context-06 for treatment-pattern distribution.",
        ],
        [
            "DME definition",
            "Publish or append the diagnosis-code/EMR logic; current public papers generally do not provide a reusable code list.",
            "This is the largest transparency gap across the corpus.",
        ],
        [
            "Follow-up",
            "For one-year VA, require a months 11-12 visit; add quarterly VA requirements only when treatment-intensity trajectory is the key endpoint.",
            "Balances cohort size and longitudinal interpretability.",
        ],
        [
            "VA method",
            "Convert Snellen to ETDRS letters with a stated formula and document accepted VA measurement methods.",
            "Core-04 and context-07 use 85 + 50*log(Snellen fraction); core-02 uses 85 - 50*logMAR.",
        ],
        [
            "Effect modifiers",
            "Prespecify baseline VA strata and injection-frequency strata.",
            "Baseline VA ceiling effects and injection intensity recur across Vestrum DME papers.",
        ],
        [
            "Sensitivity analyses",
            "Handle bilateral eyes, switchers, loss to follow-up, and competing retinal diagnoses explicitly.",
            "Core-04 uses bilateral/switcher sensitivity; context-06 excludes competing retinal pathologies.",
        ],
    ]
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["Design field", "Recommended operational choice", "Rationale"]):
        set_cell_text(table.rows[0].cells[i], h, color=LIME, bold=True, size=7.5)
    for vals in rows:
        cells = table.add_row().cells
        for i, val in enumerate(vals):
            set_cell_text(cells[i], val, size=7)
    style_table(table, widths=[1.25, 4.6, 4.6], font_size=7)


def add_source_table(doc, source_rows):
    table_caption(doc, "Source log and blocked-access caveats.", [("source log", SOURCE_LOG_JSON.as_uri())])
    relevant = [r for r in source_rows if r["id"] != "media-root"]
    table = doc.add_table(rows=1, cols=5)
    for i, h in enumerate(["ID", "Title", "Current URL", "Access status", "Retrieval"]):
        set_cell_text(table.rows[0].cells[i], h, color=LIME, bold=True, size=7.4)
    for row in relevant:
        cells = table.add_row().cells
        set_cell_text(cells[0], row["id"], size=6.6)
        set_cell_text(cells[1], row["title"], size=6.6)
        if row["current_access_url"]:
            add_doc_hyperlink_cell(cells[2], "open", row["current_access_url"])
        else:
            set_cell_text(cells[2], "No current source URL pursued", size=6.6)
        set_cell_text(cells[3], row["access_status"], size=6.3)
        set_cell_text(cells[4], row["retrieval_date"], size=6.6)
    style_table(table, widths=[0.55, 2.45, 0.8, 5.3, 0.75], font_size=6.3)


def add_landscape_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section(section, landscape=True)
    return section


def add_portrait_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section(section, landscape=False)
    return section


def build():
    rows, source_rows = load_data()
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    set_section(doc.sections[0], landscape=False)

    add_eyebrow(doc, "Retina RWE  ·  Internal Methods Reference")
    add_title(doc, "Vestrum DME Methods Comparison")
    add_body(
        doc,
        "A source-backed methods report comparing publications that used Vestrum data to evaluate diabetic macular edema, with emphasis on cohort construction, operational outcome definitions, available code lists, and practical implications for future DME studies.",
        [("Vestrum media", "https://www.vestrumhealth.com/media.php"), ("detailed extraction", DETAILED_JSON.as_uri())],
    )
    add_callout(
        doc,
        "Key finding",
        "Across the reviewed Vestrum DME corpus, public papers generally do not disclose reusable DME ICD code lists or NDC drug dictionaries. The articles instead describe diagnosis, medication, procedure, and treatment fields from the Vestrum EMR, while the most explicit code-level details appear in the cost-modeling papers.",
        [("detailed extraction", DETAILED_JSON.as_uri()), ("Grewal cost analysis", "https://pmc.ncbi.nlm.nih.gov/articles/PMC11556346/")],
    )

    add_heading(doc, "1", "Executive Summary")
    add_bullet(
        doc,
        "The core Vestrum website set contains four qualifying DME database analyses and one DME-tagged commentary that was logged but excluded from methods extraction.",
        [("Vestrum media", "https://www.vestrumhealth.com/media.php"), ("source log", SOURCE_LOG_JSON.as_uri())],
    )
    add_bullet(
        doc,
        "The most reproducible one-year anti-VEGF outcomes design is treatment-naive DME/DMO, index at first anti-VEGF treatment, one-year follow-up in months 11-12, baseline VA stratification, injection-frequency stratification, and bilateral/switcher sensitivity analysis.",
        [("BJO 28,658 eyes", "https://pmc.ncbi.nlm.nih.gov/articles/PMC7848066/")],
    )
    add_bullet(
        doc,
        "The strongest treatment-pattern design is diagnosis-indexed newly diagnosed DME with treatment categories for untreated eyes, anti-VEGF, steroid, focal laser, and combinations.",
        [("Sodhi", "https://pmc.ncbi.nlm.nih.gov/articles/PMC10170625/")],
    )
    add_bullet(
        doc,
        "Baseline VA and treatment intensity are recurring modifiers of real-world VA outcomes; worse baseline VA tends to produce larger gains, while lower injection intensity is associated with smaller gains.",
        [("BJO 28,658 eyes", "https://pmc.ncbi.nlm.nih.gov/articles/PMC7848066/"), ("longer-term outcomes", "https://www.sciencedirect.com/science/article/pii/S2468653022001506")],
    )
    add_bullet(
        doc,
        "Economic studies show that Vestrum can support utilization and HEOR analyses using visits, imaging, injections, drug mix, VA change, cost assumptions, and sensitivity analyses.",
        [("Grewal", "https://pmc.ncbi.nlm.nih.gov/articles/PMC11556346/"), ("Leung", "https://www.asrs.org/content/documents/leung-et-al-2025-cost-effectiveness-of-treatments-for-diabetic-macular-edema-simulated-bevacizumab-first-step-therapy.pdf")],
    )

    add_heading(doc, "2", "Interpretation for Future DME Analyses")
    add_body(
        doc,
        "For an anti-VEGF effectiveness analysis, the preferred starting point is a treatment-indexed cohort. The BJO 28,658-eye analysis has the broadest one-year template because it removed the loading-dose requirement used by the 2018 paper and retained a large sample for injection-frequency and baseline-VA stratification.",
        [("BJO 28,658 eyes", "https://pmc.ncbi.nlm.nih.gov/articles/PMC7848066/"), ("2018 outcomes", "https://pubmed.ncbi.nlm.nih.gov/31047187/")],
    )
    add_body(
        doc,
        "For a treatment-pattern or care-pathway analysis, a diagnosis-indexed cohort is more appropriate because it captures untreated eyes and non-anti-VEGF pathways. The Sodhi paper is the clearest example, but it also shows a limitation: ICD coding could not distinguish clinically significant from center-involving DME and imaging data were not available to verify DME severity.",
        [("Sodhi", "https://pmc.ncbi.nlm.nih.gov/articles/PMC10170625/")],
    )
    add_body(
        doc,
        "For HEOR questions, Vestrum is most useful as a utilization and outcomes input rather than as a stand-alone economic model. The Grewal and Leung analyses use Vestrum real-world visits, imaging, injections, drug mix, and VA outcomes, then layer in CMS, Medicare ASP, societal-cost, utility, and sensitivity-analysis assumptions.",
        [("Grewal", "https://pmc.ncbi.nlm.nih.gov/articles/PMC11556346/"), ("Leung", "https://www.asrs.org/content/documents/leung-et-al-2025-cost-effectiveness-of-treatments-for-diabetic-macular-edema-simulated-bevacizumab-first-step-therapy.pdf")],
    )

    add_heading(doc, "3", "Access and Evidence Caveats")
    add_body(
        doc,
        "Several older Vestrum PDF URLs now redirect away from the original files, so the local evidence folder preserves mirrors, PMC pages, PubMed pages, and access-status notes. The 2026 faricimab paper is currently abstract/indexed-page level for this report because SAGE full text and supplemental DOCX files were blocked locally, and PMC indicates full text availability on 2027-04-09.",
        [("source log", SOURCE_LOG_JSON.as_uri()), ("faricimab PubMed", "https://pubmed.ncbi.nlm.nih.gov/41971251/")],
    )

    add_landscape_section(doc)
    add_heading(doc, "4", "Publication Inventory")
    add_inventory_table(doc, rows)

    add_landscape_section(doc)
    add_heading(doc, "5", "Detailed Operational Definitions")
    add_methods_table(doc, rows)

    add_landscape_section(doc)
    add_heading(doc, "6", "Recommended Study Template")
    add_recommendation_table(doc)

    add_landscape_section(doc)
    add_heading(doc, "7", "Source Log")
    add_source_table(doc, source_rows)

    add_portrait_section(doc)
    add_heading(doc, "8", "Closing Notes")
    add_body(
        doc,
        "The detailed CSV and JSON files are the authoritative structured extraction behind this report. They retain article-level provenance, access status, and whether a field was available from full text, abstract, appendix, supplement listing, or local source log.",
        [("detailed CSV", (ROOT / "methods_extraction_detailed.csv").as_uri()), ("detailed JSON", DETAILED_JSON.as_uri())],
    )
    add_body(
        doc,
        "The central methodological improvement for future Vestrum DME publications would be to publish reusable diagnosis and treatment-code logic, or to point readers to an appendix when the list is too long for the article body.",
        [("detailed extraction", DETAILED_JSON.as_uri())],
    )

    doc.core_properties.title = "Vestrum DME Methods Comparison"
    doc.core_properties.subject = "Vestrum Health diabetic macular edema real-world evidence publication methods"
    doc.core_properties.author = "OpenAI Codex"
    doc.save(str(OUTPUT))
    clean_unused_hyperlink_relationships(OUTPUT)
    print(OUTPUT)


def clean_unused_hyperlink_relationships(path):
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    doc_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    rid_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    ET.register_namespace("", rel_ns)
    ET.register_namespace("w", doc_ns)
    ET.register_namespace("r", rid_ns)

    with ZipFile(path, "r") as zin:
        document_xml = zin.read("word/document.xml")
        rels_xml = zin.read("word/_rels/document.xml.rels")
        doc_root = ET.fromstring(document_xml)
        used = {
            el.get(f"{{{rid_ns}}}id")
            for el in doc_root.iter(f"{{{doc_ns}}}hyperlink")
            if el.get(f"{{{rid_ns}}}id")
        }
        rel_root = ET.fromstring(rels_xml)
        for rel in list(rel_root):
            if rel.get("Type", "").endswith("/hyperlink") and rel.get("Id") not in used:
                rel_root.remove(rel)
        new_rels = ET.tostring(rel_root, encoding="utf-8", xml_declaration=True)

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp_path = Path(tmp.name)
        try:
            with ZipFile(tmp_path, "w", ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == "word/_rels/document.xml.rels":
                        data = new_rels
                    zout.writestr(item, data)
            shutil.move(str(tmp_path), str(path))
        finally:
            if tmp_path.exists():
                tmp_path.unlink()


if __name__ == "__main__":
    build()
