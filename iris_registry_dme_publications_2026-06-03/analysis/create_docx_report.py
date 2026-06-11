from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REFERENCE_DOCX = Path(r"C:\Users\Justin\Desktop\LinkedIn\Target_Product_Profile_cohere_report.docx")
OUT = ROOT / "analysis" / "IRIS_Registry_DME_Methods_Comparison_cohere_style.docx"

INK = "10120F"
GRID = "1B1F17"
MUTED = "5C6257"
BLUE = "2E74B5"
TAN = "F6F1E8"
TAN_2 = "EBE4D6"
PALE = "FBF8F1"
WHITE = "FFFFFF"


SOURCES = {
    "AAO IRIS Research": "https://www.aao.org/iris-registry/research",
    "AAO Annual Meeting": "https://www.aao.org/iris-registry/annual-meeting",
    "Cantrell 2020": "https://doi.org/10.1016/j.ophtha.2019.10.019",
    "Malhotra 2021": "https://doi.org/10.1016/j.ophtha.2021.03.010",
    "Greenlee 2022": "https://doi.org/10.3928/23258160-20220615-01",
    "Maturi 2024": "https://pmc.ncbi.nlm.nih.gov/articles/PMC11102718/",
    "Maturi supplement": "https://journals.sagepub.com/doi/suppl/10.1177/24741264231221607",
    "Kuo 2024": "https://www.sciencedirect.com/science/article/pii/S2468653024002665",
    "Singh 2024": "https://pmc.ncbi.nlm.nih.gov/articles/PMC11684133/",
    "Borkar 2025": "https://pubmed.ncbi.nlm.nih.gov/40371971/",
    "Singh 2025": "https://pubmed.ncbi.nlm.nih.gov/41186474/",
    "Gong 2021": "https://pmc.ncbi.nlm.nih.gov/articles/PMC9560578/",
    "Ambrosino 2025": "https://pmc.ncbi.nlm.nih.gov/articles/PMC11999678/",
    "Zhang 2026": "https://link.springer.com/article/10.1007/s40123-026-01371-8",
    "Zhang supplement": "https://static-content.springer.com/esm/art%3A10.1007%2Fs40123-026-01371-8/MediaObjects/40123_2026_1371_MOESM1_ESM.pdf",
    "ASRS 2022 Leng": "https://medically.gene.com/content/dam/pdmahub/restricted/ophthalmology/asrs-2022/ASRS-2022-presentation-leng-long-term-real-world-treatment-patterns.pdf",
    "ARVO 2022 Kim": "https://medically.gene.com/content/dam/pdmahub/restricted/ophthalmology/arvo-2022/ARVO-2022-poster-kim-discontinuation-switching-and-other-long-term-routine-clinical-practice.pdf",
    "ASRS 2023 FARETINA": "https://medically.gene.com/content/dam/pdmahub/restricted/ophthalmology/asrs-2023/ASRS-2023-presentation-borkar-early-treatment-patterns-and-outcomes.pdf",
    "ARVO 2024 FARETINA": "https://veranahealth.com/wp-content/uploads/2024/05/ARVO-2024_FARETINA-DME-Borkar-poster_FINAL-2.pdf",
    "ASRS 2024 FARETINA": "https://medically.gene.com/content/dam/pdmahub/restricted/ophthalmology/asrs-2024/ASRS-2024-presentation-leng-12-month-real-world-clinical.pdf",
    "Hawaiian Eye 2025 FARETINA": "https://medically.gene.com/content/dam/pdmahub/restricted/ophthalmology/hawaiian-eye-2025/Hawaiian-Eye-2025-presentation-borkar-12-month-real-world-clinical-outcomes.pdf",
    "Macula Society 2026 FARETINA": "https://medically.gene.com/content/dam/pdmahub/restricted/ophthalmology/macula-society-2026/Macula-Society-2026-presentation-borkar-two-year-real-world-clinical-outcomes-in-patients.pdf",
    "ARVO 2026 linked claims": "https://medically.gene.com/content/dam/pdmahub/restricted/ophthalmology/arvo-2026/ARVO-2026-poster-cooper-real-world-outcomes-in-individuals-with-diabetic-macular-edema.pdf",
    "ISPOR 2026": "https://www.ispor.org/heor-resources/presentations-database/presentation-cti/ispor-2026/poster-session-2-4/clinical-and-economic-impact-of-faricimab-treated-diabetic-macular-edema-a-linked-analysis-of-administrative-claims-and-iris-registry-data",
}


def clear_template_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def ensure_tan_document_background(doc: Document) -> None:
    document = doc._element
    background = document.find(qn("w:background"))
    if background is None:
        background = OxmlElement("w:background")
        document.insert(0, background)
    background.set(qn("w:color"), TAN)

    settings = doc.settings._element
    if settings.find(qn("w:displayBackgroundShape")) is None:
        settings.append(OxmlElement("w:displayBackgroundShape"))


def set_section(section, landscape=False):
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, left=140, bottom=120, right=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_props(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4" if edge in ("top", "left", "bottom", "right") else "2")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), GRID)

    cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for m in ("left", "right"):
        node = cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            cell_mar.append(node)
        node.set(qn("w:w"), "10")
        node.set(qn("w:type"), "dxa")

    look = tbl_pr.find(qn("w:tblLook"))
    if look is not None:
        look.set(qn("w:val"), "0000")
        look.set(qn("w:firstRow"), "0")
        look.set(qn("w:lastRow"), "0")
        look.set(qn("w:firstColumn"), "0")
        look.set(qn("w:lastColumn"), "0")
        look.set(qn("w:noHBand"), "0")
        look.set(qn("w:noVBand"), "0")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def add_run_spacing_and_highlight(run, spacing=30, fill="D7FF5F"):
    r_pr = run._r.get_or_add_rPr()
    spacing_el = r_pr.find(qn("w:spacing"))
    if spacing_el is None:
        spacing_el = OxmlElement("w:spacing")
        r_pr.append(spacing_el)
    spacing_el.set(qn("w:val"), str(spacing))
    shd = r_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        r_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=INK, size=8.3, caps=False, line_spacing=1.1):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = line_spacing
    r = p.add_run(text)
    r.bold = bold
    r.font.all_caps = caps
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_hyperlink(paragraph, text, url, size=None, color_value="0563C1"):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_value)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    if size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))
        r_pr.append(sz)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_citations(paragraph, labels, size=9, color_value="0563C1"):
    if not labels:
        return
    paragraph.add_run(" ")
    for i, label in enumerate(labels):
        if i:
            paragraph.add_run(" ")
        add_hyperlink(paragraph, f"[{label}]", SOURCES[label], size=size, color_value=color_value)


def add_para(doc, text="", citations=None, italic=False, bold=False, size=10, color=INK, before=0, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    add_citations(p, citations or [], size=max(size - 1, 8))
    return p


def add_heading(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.333
    p_pr = p._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), INK)
    run = p.add_run(f"{number}.  {text}")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(INK)
    return p


def add_note_box(doc, text, citations=None):
    table = doc.add_table(rows=1, cols=1)
    set_table_props(table, 10080)
    cell = table.cell(0, 0)
    set_cell_shading(cell, INK)
    set_cell_margins(cell, top=240, left=280, bottom=240, right=280)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.333
    if text.startswith("Citation rule:"):
        label, rest = "Citation rule", text[len("Citation rule"):]
        r = p.add_run(label)
        r.bold = True
        r.font.all_caps = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(INK)
        add_run_spacing_and_highlight(r)
        r2 = p.add_run(rest)
        r2.font.color.rgb = RGBColor.from_string(TAN)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.color.rgb = RGBColor.from_string(TAN)
        r.font.size = Pt(10)
        r.bold = True
    add_citations(p, citations or [], size=8.5, color_value="D7FF5F")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_bullets(doc, bullets):
    for text, citations in bullets:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.28))
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run("-\t")
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string(INK)
        r2 = p.add_run(text)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = RGBColor.from_string(INK)
        add_citations(p, citations, size=8.5)


def add_table(doc, headers, rows, widths=None, font_size=7.3, table_width=14256):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True
    set_table_props(table, table_width)
    hdr = table.rows[0].cells
    prevent_row_split(table.rows[0])
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], INK)
        set_cell_margins(hdr[i], top=120, left=140, bottom=120, right=140)
        set_cell_text(hdr[i], h, bold=True, color="D7FF5F", size=font_size, caps=True, line_spacing=1.0)
        if widths:
            hdr[i].width = Inches(widths[i])
    for ri, row in enumerate(rows):
        row_obj = table.add_row()
        prevent_row_split(row_obj)
        cells = row_obj.cells
        fill = TAN if ri % 2 == 0 else TAN_2
        for ci, val in enumerate(row):
            set_cell_shading(cells[ci], fill)
            set_cell_margins(cells[ci], top=120, left=140, bottom=120, right=140)
            if widths:
                cells[ci].width = Inches(widths[ci])
            if isinstance(val, tuple):
                text, links = val
            else:
                text, links = val, []
            set_cell_text(cells[ci], text, bold=(ci == 0), size=font_size, line_spacing=1.1)
            if links:
                p = cells[ci].paragraphs[0]
                p.add_run(" ")
                for li, label in enumerate(links):
                    if li:
                        p.add_run(" ")
                    add_hyperlink(p, f"[{label}]", SOURCES[label], size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_source_links(doc):
    add_heading(doc, "7", "Source Links")
    add_para(
        doc,
        "The links below are the primary web sources and supplement locations used for the inline citations above; saved local copies and extraction notes remain in the IRIS Registry research folder.",
        ["AAO IRIS Research", "AAO Annual Meeting"],
    )
    rows = []
    for label in [
        "AAO IRIS Research",
        "AAO Annual Meeting",
        "Cantrell 2020",
        "Malhotra 2021",
        "Greenlee 2022",
        "Maturi 2024",
        "Maturi supplement",
        "Kuo 2024",
        "Singh 2024",
        "Borkar 2025",
        "Singh 2025",
        "Gong 2021",
        "Ambrosino 2025",
        "Zhang 2026",
        "Zhang supplement",
        "ARVO 2024 FARETINA",
        "ASRS 2024 FARETINA",
        "Hawaiian Eye 2025 FARETINA",
        "Macula Society 2026 FARETINA",
        "ARVO 2026 linked claims",
        "ISPOR 2026",
    ]:
        rows.append((label, ("Open source", [label])))
    add_table(doc, ["Citation label", "Hyperlink"], rows, widths=[2.2, 4.8], font_size=8)


def build_report():
    doc = Document(REFERENCE_DOCX)
    clear_template_body(doc)
    ensure_tan_document_background(doc)
    set_section(doc.sections[0], landscape=False)

    # Title block, modeled after the provided Cohere-style DOCX.
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Registry Methods  ·  Internal Evidence Review")
    r.bold = True
    r.font.all_caps = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(INK)
    add_run_spacing_and_highlight(r)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.333
    r = p.add_run("IRIS Registry DME Methods Comparison")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(INK)

    add_para(
        doc,
        "A sourced Word report summarizing how prior IRIS Registry publications, posters, presentations, and abstracts have evaluated patients with diabetic macular edema, with emphasis on patient definitions, index rules, lookback windows, outcomes, and code/supplement provenance.",
        ["AAO IRIS Research", "AAO Annual Meeting"],
        italic=True,
        size=10,
        color=MUTED,
        after=10,
    )
    add_para(
        doc,
        "Prepared from the locally saved IRIS Registry DME corpus and the expanded methods extraction dated June 3, 2026; full text was reviewed where accessible, and publisher-blocked or supplement-only details are labeled as such.",
        ["Maturi 2024", "Singh 2024", "Kuo 2024", "Zhang 2026"],
        size=9.5,
        after=14,
    )
    add_note_box(
        doc,
        "Citation rule: every finding-oriented paragraph and table row includes inline hyperlinked source labels. Long code lists are pointed to their supplement or appendix location rather than reproduced in full.",
        ["Maturi supplement", "Zhang supplement"],
    )

    add_heading(doc, "1", "Executive Takeaways")
    add_bullets(
        doc,
        [
            (
                "The most common IRIS DME study design is retrospective registry analysis using diagnosis or treatment events to define an index date, usually followed by treatment-pattern, visual-acuity, anatomic, durability, or safety outcomes.",
                ["Cantrell 2020", "Kuo 2024", "Singh 2024", "ARVO 2024 FARETINA"],
            ),
            (
                "For treatment-naive anti-VEGF studies, the most explicit operational pattern is a 12-month clean lookback with no prior anti-VEGF, a DME diagnosis near the index injection, adult age criteria, and a baseline VA window before or at index.",
                ["Singh 2024", "Kuo 2024", "ASRS 2022 Leng"],
            ),
            (
                "Visual acuity is commonly normalized to approximate ETDRS letters; the full-text studies expose the formula `85 + 50 x log(Snellen fraction)` when Snellen conversion is required.",
                ["Maturi 2024", "Singh 2024"],
            ),
            (
                "The faricimab FARETINA-DME materials add the clearest anatomic and durability definitions: treatment-naive vs prior-treated lookback, VA windows around injection visits, CST baseline/follow-up windows, interval extension thresholds, and injection-frequency summaries.",
                ["ARVO 2024 FARETINA", "ASRS 2024 FARETINA", "Hawaiian Eye 2025 FARETINA", "Macula Society 2026 FARETINA"],
            ),
            (
                "Code-level reproducibility remains uneven: some full-text papers point to supplements for ICD-10-CM or disease-code definitions, while older publisher-gated papers expose only abstract-level methods in the available public record.",
                ["Maturi supplement", "Malhotra 2021", "Greenlee 2022", "Zhang supplement"],
            ),
        ],
    )

    add_heading(doc, "2", "Recommended Methods Pattern")
    add_para(
        doc,
        "For a new IRIS Registry DME evaluation, the most reproducible approach is an eye-level retrospective cohort with laterality-specific DME identification, a clinically meaningful index date such as first anti-VEGF or first faricimab injection, a 12-month pre-index lookback to classify treatment history, and explicit follow-up windows for VA, CST, injections, discontinuation, reinitiation, switching, and safety outcomes.",
        ["Singh 2024", "Kuo 2024", "ARVO 2024 FARETINA", "ASRS 2024 FARETINA"],
    )
    add_para(
        doc,
        "If both eyes are eligible, prior studies either analyze eyes while accounting for bilateral clustering or select one eye per patient; the Maturi full text explicitly selected one eye per patient and randomly selected laterality when both eyes met criteria.",
        ["Maturi 2024", "Singh 2024"],
    )
    add_para(
        doc,
        "Outcomes that have already been studied include initial treatment choice, anti-VEGF injection use, VA change, DR severity at initiation, annualized injections, injection intervals, discontinuation, reinitiation, switching, CST change, CST threshold achievement, safety events, infectious endophthalmitis outcomes, health care resource use, and costs.",
        ["Cantrell 2020", "Malhotra 2021", "Greenlee 2022", "Kuo 2024", "Singh 2024", "Zhang 2026", "ARVO 2026 linked claims"],
    )

    # Landscape section for major comparison tables.
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section(sec, landscape=True)
    add_heading(doc, "3", "Primary DME Publication Methods Matrix")
    add_table(
        doc,
        ["Publication / item", "Design and period", "Cohort / patient definition", "Index and lookback", "Outcomes studied", "Operational details and gaps", "Source"],
        [
            [
                "Cantrell 2020 treatment patterns",
                "Retrospective IRIS Registry analysis; newly diagnosed / incident DME.",
                "13,410 treatment-naive patients with DME in accessible article summaries.",
                "Initial DME diagnosis; treatment captured within 28 days.",
                "Observation/no immediate treatment, anti-VEGF, laser, corticosteroid, and combination therapy.",
                "Initial 28-day counts: 9,990 observation/no treatment; 2,086 anti-VEGF; 1,133 laser; 133 corticosteroid; 68 combination. Code list and full inclusion/exclusion not exposed in accessible source.",
                ("", ["Cantrell 2020", "AAO IRIS Research"]),
            ],
            [
                "Malhotra 2021 disparities at anti-VEGF initiation",
                "Retrospective cross-sectional IRIS Registry study, 2012-2020.",
                "Patients initiating anti-VEGF injection treatment for DME; n about 203,707 in abstract/preview.",
                "Anti-VEGF initiation for DME.",
                "Baseline/presenting VA and DR severity at treatment initiation.",
                "Compared race, ethnicity, insurance, and geography using multivariable regression. Diagnostic code set and VA cleaning/window rules require publisher supplement/full text.",
                ("", ["Malhotra 2021"]),
            ],
            [
                "Greenlee 2022 socioeconomic disparities",
                "Retrospective IRIS Registry cohort, 2012-2020.",
                "DME diagnosis plus at least one anti-VEGF injection; n 203,707 in PubMed abstract.",
                "Anti-VEGF-treated DME cohort; exact index window not exposed in public abstract.",
                "Anti-VEGF injection use over 60 months and longitudinal VA outcomes.",
                "PubMed-level methods report multivariate regression, incidence rate ratios for injection use, and odds ratios for longitudinal VA; exact code lists and VA algorithms not exposed.",
                ("", ["Greenlee 2022"]),
            ],
            [
                "Maturi 2024 race/insurance in DR/DME",
                "Retrospective IRIS Registry analysis; query January 1, 2014 to December 31, 2018.",
                "Adults with DR, anti-VEGF treatment, and baseline/1-year/2-year VA; final cohort 43,274 eyes.",
                "Anti-VEGF treatment after DR diagnosis; one eye per patient, randomly selected if both eligible.",
                "Mean VA change at 1 and 2 years; bevacizumab use; DR severity; DME presence; >=15-letter loss in supplemental figures.",
                "ICD-10-CM used for DR severity and DME presence; code definitions in supplement. ETDRS conversion formula exposed. Combination therapy = >1 anti-VEGF drug in first 365 days.",
                ("", ["Maturi 2024", "Maturi supplement"]),
            ],
            [
                "Kuo 2024 long-term treatment patterns",
                "Retrospective IRIS Registry analysis; treatment-naive DME eyes initiating anti-VEGF IVT from January 1, 2015 to March 31, 2021.",
                "190,345 eyes in abstract/preview.",
                "First anti-VEGF IVT after no prior IVT in previous 12 months; follow-up up to 6 years.",
                "Annualized injections, injection intervals, anti-VEGF agent use, VA change, discontinuation, switching, and reinitiation.",
                "Discontinuation defined in preview as >1 year / 365 days contributing to IRIS without anti-VEGF IVT. Full code definitions and supplemental tables require Ophthalmology Retina supplement.",
                ("", ["Kuo 2024", "ASRS 2022 Leng", "ARVO 2022 Kim"]),
            ],
            [
                "Singh 2024 initial anti-VEGF dosing",
                "Retrospective IRIS Registry database study; index January 1, 2015 to December 31, 2020; data cutoff December 31, 2021.",
                "Adults with DME within 2 months of index, >=12 months pre-index data, baseline VA at/within 60 days before index, anti-VEGF at index, and >=100 days follow-up; final 217,696 eyes.",
                "Index = earliest anti-VEGF injection with no anti-VEGF in prior 12 months; initial-dose = >=3 same-agent injections within 100 days.",
                "VA change, discontinuation, reinitiation, switching, injection patterns, and initial-dose likelihood.",
                "Discontinuation = >365 days contributing to IRIS without another anti-VEGF after last injection; reinitiation = anti-VEGF >=1 year after last injection; ETDRS conversion formula exposed.",
                ("", ["Singh 2024"]),
            ],
            [
                "Borkar 2025 / Singh 2025 FARETINA-DME manuscripts",
                "IRIS Registry FARETINA-DME faricimab manuscripts; PubMed/abstract-level methods available publicly.",
                "DME patients initiating faricimab; manuscripts report treatment-naive and previously anti-VEGF-treated cohorts.",
                "Index = faricimab initiation; detailed windows best exposed in conference materials.",
                "VA at injection 4 or injection 7, CST change, CST threshold achievement, injection frequency, and durability.",
                "Public manuscripts are abstract-level locally; FARETINA conference PDFs expose treatment-naive lookback, VA/CST windows, and interval-extension definitions.",
                ("", ["Borkar 2025", "Singh 2025", "ARVO 2024 FARETINA", "ASRS 2024 FARETINA"]),
            ],
        ],
        widths=[1.35, 1.45, 1.7, 1.55, 1.65, 2.6, 1.2],
        font_size=6.8,
    )

    add_heading(doc, "4", "Operational Definitions and Outcome Algorithms")
    add_table(
        doc,
        ["Domain", "Definitions observed in the IRIS DME corpus", "Best source / provenance"],
        [
            [
                "DME / DR identification",
                "DME is generally captured from IRIS diagnosis fields; Maturi explicitly used ICD-10-CM codes for DR severity and DME presence, and Zhang places DME/DR disease codes in a long supplementary table.",
                ("Maturi code definitions are in the SAGE/PMC supplement; Zhang disease-code lists are in Supplementary file 1, Table S1.", ["Maturi 2024", "Maturi supplement", "Zhang supplement"]),
            ],
            [
                "Treatment-naive anti-VEGF",
                "The clearest pattern is no anti-VEGF or IVT anti-VEGF during a >=12-month pre-index lookback before the first qualifying injection.",
                ("Used by Singh 2024 and reflected in Kuo/ASRS long-term treatment pattern materials.", ["Singh 2024", "Kuo 2024", "ASRS 2022 Leng"]),
            ],
            [
                "Initial-dose anti-VEGF",
                "Initial-dose cohort = >=3 injections of the same anti-VEGF agent within the first 100 days after the index injection; non-initial cohort = otherwise eligible eyes not meeting that rule.",
                ("Definition from the Singh 2024 full text.", ["Singh 2024"]),
            ],
            [
                "Discontinuation, reinitiation, switching",
                "Discontinuation commonly uses a >365-day / >1-year period contributing to IRIS with no subsequent anti-VEGF injection. Reinitiation is anti-VEGF after >=1 year, and switching is a different anti-VEGF agent after the defined gap; conference precursor materials also used >=3 consecutive injections of a different anti-VEGF agent for switching.",
                ("Definitions are exposed most clearly in Singh 2024 and long-term treatment pattern conference materials.", ["Singh 2024", "Kuo 2024", "ARVO 2022 Kim", "ASRS 2022 Leng"]),
            ],
            [
                "Visual acuity",
                "VA is typically converted to approximate ETDRS letters. The exposed formula is `85 + 50 x log(Snellen fraction)`. Singh used annual best VA; FARETINA materials include corrected, uncorrected, and pinhole VA around injection visits.",
                ("Formula and VA handling from Maturi and Singh; FARETINA windows from posters/presentations.", ["Maturi 2024", "Singh 2024", "ARVO 2024 FARETINA"]),
            ],
            [
                "CST / anatomy",
                "FARETINA-DME uses baseline CST 0-30 days before index in extracted materials, requires repeated CST measures before and after index in some analyses, excludes CST values shortly after injections in longitudinal windows, and reports mean CST change plus CST <=280 micrometers or >=10% reduction.",
                ("CST availability is limited and window definitions are from FARETINA conference PDFs.", ["ARVO 2024 FARETINA", "ASRS 2024 FARETINA", "Macula Society 2026 FARETINA"]),
            ],
            [
                "Injection interval / durability",
                "FARETINA materials classify interval extension after repeated faricimab injections; an extended interval is >6 weeks after the previous faricimab injection, with interval bands such as <8, >=8 to <12, and >=12 weeks.",
                ("Definition from FARETINA conference materials.", ["ASRS 2024 FARETINA", "Hawaiian Eye 2025 FARETINA", "Macula Society 2026 FARETINA"]),
            ],
            [
                "Ocular safety events",
                "FARETINA safety definitions use incident ICD-10 diagnoses after faricimab initiation with no corresponding diagnosis in the prior 12 months; extracted code groups include endophthalmitis, iridocyclitis/iritis, retinal vasculitis, uveitis, and vitritis.",
                ("Specific ICD-10 codes are in the detailed methods extraction and conference PDFs; code-based incidence should be interpreted cautiously.", ["ARVO 2024 FARETINA", "ASRS 2024 FARETINA"]),
            ],
            [
                "Infectious endophthalmitis outcomes",
                "Zhang identified IVT biologic exposure using HCPCS J-codes; infectious endophthalmitis codes include ICD-9 360.00, 360.01, 360.03, 360.19 and ICD-10 H44.0, H44.001, H44.002, H44.003, H44.19. Legal blindness is <=35 letters or <=20/200.",
                ("Disease subgroup code lists are too long to reproduce; use Supplementary file 1, Table S1.", ["Zhang 2026", "Zhang supplement"]),
            ],
            [
                "Linked claims / economic outcomes",
                "The ARVO 2026 linked analysis combines CVS administrative claims with IRIS EHR data, requires DME claims and continuous medical/pharmacy coverage around faricimab initiation, and compares VA/CST, injections, visits, costs, and health care resource utilization.",
                ("Operational detail from the ARVO 2026 poster; ISPOR abstract overlaps at a higher level.", ["ARVO 2026 linked claims", "ISPOR 2026"]),
            ],
        ],
        widths=[1.6, 5.1, 3.2],
        font_size=7.2,
    )

    add_heading(doc, "5", "Secondary and Access-Limited Items")
    add_table(
        doc,
        ["Item", "Why it matters", "Specific methods detail retained", "Access / citation note"],
        [
            [
                "AAO 2017 PA017 presentation listing",
                "AAO-listed precursor to the treatment-pattern topic.",
                "Purpose was characterization of treatment patterns surrounding incident DME in the United States; no public presentation file was exposed.",
                ("Use Cantrell 2020 for published detail unless the AAO archive file is obtained.", ["AAO Annual Meeting", "Cantrell 2020"]),
            ],
            [
                "Gong 2021 PDR treatment trends",
                "DME appears as a status variable in a PDR treatment-pattern analysis.",
                "Newly diagnosed PDR cohort of 141,317 patients; PDR codes explicitly reported as ICD-10 E08.35, E09.35, E10.35, E11.35, E13.35 and ICD-9 362.02.",
                ("Exact DME status codes and IVI/PRP procedure code lists were not exposed in the accessible abstract/metadata text.", ["Gong 2021"]),
            ],
            [
                "Ambrosino 2025 sickle cell trait/disease",
                "DME is an outcome/complication among patients with diabetes and sickle cell status.",
                "Cross-sectional IRIS analysis of type 1/type 2 diabetes from 2013-2021; DME, PDR, procedures, VA categories, and multivariable models were extracted.",
                ("Not a primary DME treatment-pattern study; useful for DME-as-outcome methods.", ["Ambrosino 2025"]),
            ],
            [
                "Zhang 2026 IVT biologic endophthalmitis",
                "DME/DR is a disease indication subgroup for safety outcomes after intravitreal biologics.",
                "Defines infectious endophthalmitis and evisceration/enucleation codes, legal blindness threshold, disease-stratification exclusion for overlapping disease codes, and supplement location for long disease-code lists.",
                ("Use Springer full text and Supplementary file 1, Table S1 for disease codes.", ["Zhang 2026", "Zhang supplement"]),
            ],
            [
                "Publisher-gated older DME disparity papers",
                "They are central to DME disparities but exposed public methods are incomplete.",
                "Public abstracts/previews capture broad cohort, exposures, outcomes, and statistical approach, but not code lists, VA cleaning, or detailed inclusion/exclusion.",
                ("Supplement/full text should be retrieved through institutional access for exact operational replication.", ["Malhotra 2021", "Greenlee 2022"]),
            ],
        ],
        widths=[2.0, 2.3, 3.1, 2.2],
        font_size=7.2,
    )

    # Back to portrait for synthesis and source links.
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section(sec, landscape=False)
    add_heading(doc, "6", "Interpretation for Future IRIS DME Work")
    add_para(
        doc,
        "The strongest template for a future DME outcomes study is to combine Singh-style cohort rigor with FARETINA-style outcome windows: define the index injection, require at least 12 months of pre-index record availability, classify treatment history from the lookback, require a baseline VA close to index, and pre-specify injection, VA, CST, and durability windows.",
        ["Singh 2024", "ARVO 2024 FARETINA", "ASRS 2024 FARETINA"],
    )
    add_para(
        doc,
        "For safety or disease-subgroup work, code provenance should be elevated into the main methods section or a retained supplement, because Zhang demonstrates that disease-code tables may be too long for narrative text while still being essential for replication.",
        ["Zhang 2026", "Zhang supplement"],
    )
    add_para(
        doc,
        "For disparity or treatment-access questions, prior IRIS DME studies have compared race, ethnicity, insurance, and geography against VA, DR severity, injection use, and longitudinal VA outcomes, but exact operational replication may require publisher supplements or full-text access for older Ophthalmology/Healio manuscripts.",
        ["Malhotra 2021", "Greenlee 2022", "Maturi 2024"],
    )
    add_para(
        doc,
        "For linked clinical-economic analyses, the 2026 claims-plus-IRIS work shows a path to pairing IRIS VA/CST outcomes with claims-based resource use and cost outcomes, but it also adds enrollment-continuity and linkage requirements that are not present in IRIS-only studies.",
        ["ARVO 2026 linked claims", "ISPOR 2026"],
    )

    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section(sec, landscape=True)
    add_source_links(doc)

    doc.core_properties.title = "IRIS Registry DME Methods Comparison"
    doc.core_properties.subject = "Methods comparison of IRIS Registry DME studies"
    doc.core_properties.author = "OpenAI Codex"
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_report())
