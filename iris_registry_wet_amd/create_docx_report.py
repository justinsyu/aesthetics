import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
TEMPLATE = Path(r"C:\Users\Justin\Desktop\LinkedIn\Target_Product_Profile_cohere_report.docx")
OUT = ROOT / "IRIS_Registry_wet_AMD_methods_cohere_report.docx"

DARK = "10120F"
TEXT = "1B1F17"
MUTED = "5C6257"
CREAM = "F6F1E8"
LIME = "D7FF5F"
PALE = "F3F7E8"
LINE = TEXT
TABLE_FONT_SIZE = 9.5
TABLE_TEXT_WIDTH_IN = 9.9


def clear_doc(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_margins(section, landscape=False):
    if landscape:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
    else:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def set_orientation(section, landscape=False):
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    set_margins(section, landscape)


def clear_headers_and_footers(doc):
    for section in doc.sections:
        for part in (section.header, section.first_page_header, section.even_page_header):
            part.is_linked_to_previous = False
            for paragraph in part.paragraphs:
                paragraph.clear()
        for part in (section.footer, section.first_page_footer, section.even_page_footer):
            part.is_linked_to_previous = False
            for paragraph in part.paragraphs:
                paragraph.clear()


def display_text(value):
    return str(value).replace(">=", "≥").replace("<=", "≤")


def capitalize_first(value):
    text = display_text(value).strip()
    return text[:1].upper() + text[1:] if text else text


def title_style_label(value):
    text = display_text(value).replace("/", " / ")
    words = []

    def cap_word(word):
        if word.upper() in {"AAO", "AMD", "ARVO", "FARETINA", "IRIS", "PDF"}:
            return word.upper()
        if word.lower() == "namd":
            return "nAMD"
        if "-" in word:
            return "-".join(cap_word(part) for part in word.split("-"))
        return word[:1].upper() + word[1:].lower()

    for word in text.split():
        if word == "/":
            words.append(word)
        else:
            words.append(cap_word(word))
    return " ".join(words)


def set_document_background(doc, color=CREAM):
    background = doc._element.find(qn("w:background"))
    if background is None:
        background = OxmlElement("w:background")
        doc._element.insert(0, background)
    background.set(qn("w:color"), color)


def set_run(
    run,
    size=10,
    bold=False,
    italic=False,
    color=TEXT,
    font_name="Inter",
    caps=False,
    shading=None,
    spacing=None,
):
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    if caps:
        caps_el = r_pr.find(qn("w:caps"))
        if caps_el is None:
            r_pr.append(OxmlElement("w:caps"))
    if shading:
        shd = r_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            r_pr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), shading)
    if spacing is not None:
        sp = r_pr.find(qn("w:spacing"))
        if sp is None:
            sp = OxmlElement("w:spacing")
            r_pr.append(sp)
        sp.set(qn("w:val"), str(spacing))


def add_hyperlink(paragraph, text, url, color="025E8D", bold=False, size=10.5, font_name="Inter"):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    if bold:
        b = OxmlElement("w:b")
        r_pr.append(b)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:eastAsia"), font_name)
    r_pr.append(fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz_cs)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 0
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "360")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "180")
    p_pr.append(ind)
    lvl.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Symbol")
    fonts.set(qn("w:hAnsi"), "Symbol")
    r_pr.append(fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_bullet_indent(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.insert(0, num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_el = num_pr.find(qn("w:numId"))
    if num_id_el is None:
        num_id_el = OxmlElement("w:numId")
        num_pr.append(num_id_el)
    num_id_el.set(qn("w:val"), str(num_id))

    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "180")


def add_cited_para(doc, text, citations, size=10.5, italic=False, bullet_num_id=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25
    if bullet_num_id is not None:
        apply_bullet_indent(p, bullet_num_id)
    r = p.add_run(display_text(text) + " ")
    set_run(r, size=size, italic=italic)
    for idx, (label, url) in enumerate(citations):
        add_hyperlink(p, f"[{label}]", url, size=size)
        if idx < len(citations) - 1:
            r2 = p.add_run(" ")
            set_run(r2, size=size)
    return p


def add_section_heading(doc, text, number=None):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.333
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), DARK)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    r = p.add_run(text)
    set_run(r, size=15, bold=True, color=DARK)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")


def set_table_borders(table, color=LINE):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4" if edge in ("top", "left", "bottom", "right") else "2")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_fixed_layout(table):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(int(TABLE_TEXT_WIDTH_IN * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_width(cell, width_inches):
    width_twips = int(width_inches * 1440)
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.insert(0, tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def width_profile(headers):
    profiles = {
        ("Year", "Scope", "Item", "Format", "Qualification", "Source"): [0.65, 0.85, 2.85, 1.25, 3.05, 1.25],
        ("Study", "Population / Index", "Outcome Definitions", "Analysis", "Code / Appendix Pointer", "Source"): [1.35, 1.95, 2.35, 1.45, 1.75, 1.05],
        ("Component", "Recommended operationalization", "Why it is preferred", "Source"): [1.25, 3.25, 3.15, 2.25],
        ("Citation", "Use in report", "URL"): [6.25, 2.45, 1.20],
    }
    weights = profiles.get(tuple(headers))
    if weights is None:
        weights = [1] * len(headers)
    total = sum(weights)
    return [TABLE_TEXT_WIDTH_IN * weight / total for weight in weights]


def apply_table_widths(table, widths):
    set_table_fixed_layout(table)
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)


def set_cell_margins(cell, top=90, left=110, bottom=90, right=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in {"top": top, "left": left, "bottom": bottom, "right": right}.items():
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_borders(cell, color=LINE):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def clear_cell(cell):
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
        p.text = ""


def cell_text(cell, text, bold=False, size=TABLE_FONT_SIZE, color=TEXT, hyperlink=None):
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    text = display_text(text)
    if hyperlink:
        add_hyperlink(p, text, hyperlink, color="025E8D", bold=bold, size=size)
    else:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def cell_links(cell, links, size=TABLE_FONT_SIZE):
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    for idx, (label, url) in enumerate(links):
        add_hyperlink(p, label, url, color="025E8D", size=size)
        if idx < len(links) - 1:
            spacer = p.add_run("; ")
            set_run(spacer, size=size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(doc, headers, rows, title=None):
    if title:
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(title)
        set_run(r, size=12, bold=True, color=DARK)
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Normal Table"
    except KeyError:
        pass
    table.autofit = False
    widths = width_profile(headers)
    set_table_borders(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, DARK)
        set_cell_margins(cell)
        set_cell_borders(cell)
        clear_cell(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(header)
        set_run(r, bold=True, size=TABLE_FONT_SIZE, color=LIME, caps=True, spacing=4)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, header in enumerate(headers):
            val = row.get(header, "")
            cell = cells[ci]
            shade_cell(cell, CREAM if ri % 2 == 0 else "EBE4D6")
            set_cell_margins(cell)
            set_cell_borders(cell)
            if isinstance(val, tuple):
                text, url = val
                cell_text(cell, text, size=TABLE_FONT_SIZE, hyperlink=url)
            elif isinstance(val, list):
                cell_links(cell, val, size=TABLE_FONT_SIZE)
            else:
                cell_text(cell, str(val), size=TABLE_FONT_SIZE, bold=(ci == 0))
    apply_table_widths(table, widths)
    return table


def add_callout(doc, title, body, citations):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_table_borders(table)
    shade_cell(cell, DARK)
    set_cell_margins(cell, top=130, left=150, bottom=130, right=150)
    set_cell_borders(cell, DARK)
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(title + "\n")
    set_run(r, size=TABLE_FONT_SIZE, bold=True, color=DARK, caps=True, shading=LIME, spacing=30)
    r2 = p.add_run(display_text(body) + " ")
    set_run(r2, size=9.5, color=CREAM)
    for label, url in citations:
        add_hyperlink(p, f"[{label}]", url, color=LIME, size=9.5)
        spacer = p.add_run(" ")
        set_run(spacer, size=9.5, color=CREAM)
    return table


def add_table_footer(doc, text, citations=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(display_text(text) + " ")
    set_run(r, size=8.5, italic=True, color=MUTED)
    for idx, (label, url) in enumerate(citations or []):
        add_hyperlink(p, f"[{label}]", url, color="025E8D", size=8.5)
        if idx < len(citations) - 1:
            spacer = p.add_run(" ")
            set_run(spacer, size=8.5, italic=True, color=MUTED)
    return p


def new_section(doc, landscape=False):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_orientation(section, landscape)
    return section


def best_url(item):
    urls = item.get("landing_urls") or item.get("pdf_urls") or []
    return urls[0] if urls else "https://www.aao.org/iris-registry/data-analysis/requirements"


def source_label(url):
    u = url.lower()
    if "aao.org" in u:
        return "AAO IRIS"
    if "pubmed.ncbi" in u:
        return "PubMed"
    if "pmc.ncbi" in u:
        return "PMC"
    if "springer" in u:
        return "Springer"
    if "sciencedirect" in u:
        return "ScienceDirect"
    if "asrs.org" in u:
        return "ASRS PDF"
    if "medically.gene.com" in u:
        return "ASRS 2021 PDF"
    if "retinasociety" in u:
        return "Retina Society PDF"
    if "veranahealth" in u:
        return "Verana"
    if "iovs.arvojournals" in u:
        return "IOVS/ARVO"
    if "researchsquare" in u:
        return "Research Square"
    if "doi.org" in u:
        return "DOI"
    return "Source"


def main():
    manifest = json.loads((ROOT / "publication_manifest.json").read_text(encoding="utf-8"))
    by_slug = {item["slug"]: item for item in manifest}

    AAO = "https://www.aao.org/iris-registry/data-analysis/requirements"
    MANIFEST_URL = (ROOT / "publication_manifest.md").resolve().as_uri()
    SOURCES = {
        "AAO IRIS": AAO,
        "Wykoff 2024": "https://pmc.ncbi.nlm.nih.gov/articles/PMC10767511/",
        "Gallivan 2023": "https://pmc.ncbi.nlm.nih.gov/articles/PMC10748734/",
        "Fevrier 2024": "https://pmc.ncbi.nlm.nih.gov/articles/PMC11179401/",
        "Khanani 2022": "https://pmc.ncbi.nlm.nih.gov/articles/PMC8613703/",
        "Khanani supplement": "https://pmc.ncbi.nlm.nih.gov/articles/instance/8613703/bin/jamaophthalmol-e214585-s001.pdf",
        "Zarbin 2024": "https://link.springer.com/article/10.1007/s40123-024-00920-3",
        "Zarbin ESM": "https://static-content.springer.com/esm/art%3A10.1007%2Fs40123-024-00920-3/MediaObjects/40123_2024_920_MOESM1_ESM.pdf",
        "Wykoff supplement": "https://pmc.ncbi.nlm.nih.gov/articles/instance/10767511/bin/mmc2.pdf",
        "MacCumber 2023 CJO": "https://pubmed.ncbi.nlm.nih.gov/34863677/",
        "MacCumber 2023 interval": "https://pubmed.ncbi.nlm.nih.gov/36990322/",
        "Khurana 2023": "https://pubmed.ncbi.nlm.nih.gov/36858288/",
        "Gong 2024": "https://pubmed.ncbi.nlm.nih.gov/38319061/",
        "Acharya 2025": "https://pubmed.ncbi.nlm.nih.gov/40738331/",
        "Barikian 2026": "https://pubmed.ncbi.nlm.nih.gov/40614931/",
        "Ali 2025": "https://pubmed.ncbi.nlm.nih.gov/40371970/",
        "Tabano 2026": "https://pubmed.ncbi.nlm.nih.gov/41891889/",
        "ASRS 2022": "https://www.asrs.org/content/documents/wet-amd-1.pdf",
        "Leng 2021 ASRS": "https://medically.gene.com/content/dam/pdmahub/restricted/ophthalmology/asrs-2021/ASRS-2021-presentation-leng-long-term-experience-W-ith-intravitreal-anti-VEGF-in-patients-with-nAMD-analysis-of-intelligent-research-in-sight.pdf",
        "Leng 2024": "https://veranahealth.com/wp-content/uploads/2024/05/Leng_ARVO-2024_nAMD-IRIS-Registry_Presentation_disclosure.pdf",
        "FARETINA ARVO": "https://iovs.arvojournals.org/article.aspx?articleid=2787245",
        "Khurana LTFU": "https://www.retinasociety.org/content/meetingarchive/2020/khurana-rahul-loss-to-follow-up.pdf",
        "Research Square": "https://www.researchsquare.com/article/rs-5505014/v2",
        "Saved manifest": MANIFEST_URL,
    }
    CITATION_TEXTS = {
        "AAO IRIS": "American Academy of Ophthalmology. IRIS Registry Data Analysis: Requirements. Accessed June 4, 2026. https://www.aao.org/iris-registry/data-analysis/requirements",
        "Wykoff 2024": "Wykoff CC, Garmo V, Tabano D, et al. Impact of Anti-VEGF Treatment and Patient Characteristics on Vision Outcomes in Neovascular Age-related Macular Degeneration: Up to 6-Year Analysis of the AAO IRIS Registry. Ophthalmol Sci. 2024;4(2):100421. doi:10.1016/j.xops.2023.100421",
        "Gallivan 2023": "Gallivan MD, Garcia KM, Torres AZ, et al. Emulating VIEW 1 and VIEW 2 Clinical Trial Outcome Data Using the American Academy of Ophthalmology IRIS Registry. Ophthalmic Surg Lasers Imaging Retina. 2023;54(1):6-14. doi:10.3928/23258160-20221214-01",
        "Fevrier 2024": "Fevrier H, LaPrise A, Mbagwu M, Leng T, Torres AZ, Borkar DS. Comparison of Methods of Clinical Trial Emulation Utilizing Data From the Comparison of AMD Treatment Trial (CATT) and the IRIS Registry. Ophthalmol Sci. 2024;4(5):100524. doi:10.1016/j.xops.2024.100524",
        "Khanani 2022": "Khanani AM, Zarbin MA, Barakat MR, et al. Safety Outcomes of Brolucizumab in Neovascular Age-Related Macular Degeneration: Results From the IRIS Registry and Komodo Healthcare Map. JAMA Ophthalmol. 2022;140(1):20-28. doi:10.1001/jamaophthalmol.2021.4585",
        "Khanani supplement": "Khanani AM, Zarbin MA, Barakat MR, et al. Supplement to: Safety Outcomes of Brolucizumab in Neovascular Age-Related Macular Degeneration: Results From the IRIS Registry and Komodo Healthcare Map. JAMA Ophthalmol. 2022;140(1):20-28.",
        "Zarbin 2024": "Zarbin MA, MacCumber MW, Karcher H, et al. Real-World Safety Outcomes with Brolucizumab in Neovascular Age-Related Macular Degeneration: Findings from the IRIS Registry. Ophthalmol Ther. 2024;13(5):1357-1368. doi:10.1007/s40123-024-00920-3",
        "Zarbin ESM": "Zarbin MA, MacCumber MW, Karcher H, et al. Electronic supplementary material to: Real-World Safety Outcomes with Brolucizumab in Neovascular Age-Related Macular Degeneration: Findings from the IRIS Registry. Ophthalmol Ther. 2024;13(5):1357-1368.",
        "Wykoff supplement": "Wykoff CC, Garmo V, Tabano D, et al. Supplementary Tables to: Impact of Anti-VEGF Treatment and Patient Characteristics on Vision Outcomes in Neovascular Age-related Macular Degeneration. Ophthalmol Sci. 2024;4(2):100421.",
        "MacCumber 2023 CJO": "MacCumber MW, Yu JS, Sagkriotis A, et al. Antivascular endothelial growth factor agents for wet age-related macular degeneration: an IRIS registry analysis. Can J Ophthalmol. 2023;58(3):252-261. doi:10.1016/j.jcjo.2021.10.008",
        "MacCumber 2023 interval": "MacCumber MW, Wykoff CC, Karcher H, et al. Factors Linked to Injection Interval Extension in Eyes with Wet Age-Related Macular Degeneration Switched to Brolucizumab. Ophthalmology. 2023;130(8):795-803. doi:10.1016/j.ophtha.2023.03.017",
        "Khurana 2023": "Khurana RN, Li C, Lum F. Loss to Follow-up in Patients with Neovascular Age-Related Macular Degeneration Treated with Anti-VEGF Therapy in the United States in the IRIS Registry. Ophthalmology. 2023;130(7):672-683. doi:10.1016/j.ophtha.2023.02.021",
        "Gong 2024": "Gong D, Ross C, Hall N, et al. Fellow Eyes Conversion Rates in Patients With Unilateral Exudative Age-Related Macular Degeneration: An Academy IRIS Registry Analysis. Ophthalmic Surg Lasers Imaging Retina. 2024;55(4):220-226. doi:10.3928/23258160-20240125-01",
        "Acharya 2025": "Acharya B, Momenaei B, Zhang Q, Hyman L, Haller JA, IRIS Analytic Center Consortium. Disparities in Presentation and Anti-VEGF Therapy Initiation for Neovascular Age-Related Macular Degeneration: An Analysis of the Academy IRIS Registry (Intelligent Research in Sight). Ophthalmology. 2025;132(12):1411-1421. doi:10.1016/j.ophtha.2025.07.024",
        "Barikian 2026": "Barikian A, Kumar JB, McCullough AJ, et al. Characteristics and Outcomes of Patients with Neovascular Age-Related Macular Degeneration by Anti-VEGF Exposure in United States Clinical Practice. Ophthalmol Retina. 2026;10(1):71-80. doi:10.1016/j.oret.2025.06.016",
        "Ali 2025": "Ali FS, Tabano DC, Borkar DS, et al. Early Outcomes After Initiation of Faricimab for Neovascular Age-Related Macular Degeneration. Ophthalmic Surg Lasers Imaging Retina. 2025;56(8):468-477. doi:10.3928/23258160-20250304-02",
        "Tabano 2026": "Tabano DC, Ali FS, Borkar DS, et al. One-year Real-world Outcomes With Faricimab in Neovascular Age-related Macular Degeneration. Ophthalmic Surg Lasers Imaging Retina. 2026;57(4):242-250. doi:10.3928/23258160-20260302-02",
        "ASRS 2022": "American Society of Retina Specialists. Conversion Rates from Nonexudative to Exudative Age-Related Macular Degeneration: An AAO IRIS Registry Analysis. Presented at: ASRS Annual Meeting; 2022.",
        "Leng 2021 ASRS": "Leng T, et al. Long-term Experience With Intravitreal Anti-VEGF in Patients With nAMD: Analysis of Intelligent Research in Sight Registry. Presented at: ASRS Annual Meeting; 2021.",
        "Leng 2024": "Leng T, et al. IRIS Registry Analysis of Anti-VEGF Treatment in Patients With Coexisting Neovascular Age-Related Macular Degeneration and Geographic Atrophy. Presented at: ARVO Annual Meeting; 2024.",
        "FARETINA ARVO": "Verana Health. Early Real-World Patient Outcomes After Faricimab Treatment in Neovascular Age-Related Macular Degeneration: The FARETINA-AMD Study. Invest Ophthalmol Vis Sci. 2023;64(8):abstract 5192.",
        "Khurana LTFU": "Khurana RN, et al. Loss to Follow-Up in Patients with Neovascular Age-related Macular Degeneration. Presented at: Retina Society Annual Meeting; 2020.",
        "Research Square": "Ashourizadeh H, et al. Cataract Surgery and the Risk of Conversion from Dry to Neovascular Age-related Macular Degeneration in the IRIS Registry. Research Square. Preprint posted 2026. doi:10.21203/rs.3.rs-5505014/v2",
        "Saved manifest": "IRIS Registry Wet AMD / nAMD Publication Manifest. Local research manifest. Accessed June 4, 2026.",
    }

    doc = Document(str(TEMPLATE))
    clear_doc(doc)
    set_document_background(doc)
    set_orientation(doc.sections[0], False)
    styles = doc.styles
    styles["Normal"].font.name = "Inter"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].font.color.rgb = RGBColor.from_string(TEXT)

    # Title page
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("IRIS REGISTRY  ·  WET AGE-RELATED MACULAR DEGENERATION METHODS REVIEW")
    set_run(r, size=8, bold=True, color=DARK, caps=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.333
    r = p.add_run("Evaluating Patients with Wet Age-related Macular Degeneration Using the IRIS Registry")
    set_run(r, size=22, bold=True, color=DARK, font_name="Space Grotesk")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.333
    r = p.add_run("Evidence inventory, detailed methods extraction, and recommended operational design pattern")
    set_run(r, size=11, italic=True, color=MUTED)

    add_callout(
        doc,
        "Evidence base",
        "This report synthesizes AAO IRIS Registry-listed wet AMD/nAMD publications, PubMed/PMC records, publisher full-text pages, and public conference materials saved in the local IRIS Registry research folder.",
        [("AAO IRIS", AAO), ("Saved manifest", SOURCES["Saved manifest"])],
    )

    add_section_heading(doc, "Executive Readout", 1)
    add_cited_para(
        doc,
        "Across the reviewed literature, the dominant IRIS Registry design for wet AMD/nAMD is a retrospective eye-level or patient-eye-level cohort anchored to a laterality-specific diagnosis and an anti-VEGF injection, with visual acuity and injection-pattern outcomes assessed longitudinally.",
        [("Wykoff 2024", SOURCES["Wykoff 2024"]), ("Gallivan 2023", SOURCES["Gallivan 2023"]), ("Fevrier 2024", SOURCES["Fevrier 2024"])],
    )
    add_cited_para(
        doc,
        "The most reusable operational framework is: define nAMD using ICD diagnosis codes, require same-eye anti-VEGF evidence, set index date as first qualifying injection or first drug-specific injection, require pre-index data contribution, specify VA windows, and predefine treatment gaps, discontinuation, switching, and loss-to-follow-up.",
        [("Wykoff 2024", SOURCES["Wykoff 2024"]), ("Gallivan 2023", SOURCES["Gallivan 2023"]), ("Khanani 2022", SOURCES["Khanani 2022"]), ("MacCumber interval", SOURCES["MacCumber 2023 interval"]), ("FARETINA ARVO", SOURCES["FARETINA ARVO"]), ("Khurana 2023", SOURCES["Khurana 2023"])],
    )
    add_cited_para(
        doc,
        "Previously studied outcomes include VA at diagnosis and follow-up, VA change by treatment exposure, injection frequency and intervals, switching, treatment discontinuation, LTFU/nonpersistence, brolucizumab safety, faricimab outcomes, fellow-eye conversion, and trial emulation against VIEW and CATT.",
        [("AAO IRIS", AAO), ("Wykoff 2024", SOURCES["Wykoff 2024"]), ("MacCumber CJO", SOURCES["MacCumber 2023 CJO"]), ("Khurana 2023", SOURCES["Khurana 2023"]), ("Khanani 2022", SOURCES["Khanani 2022"]), ("FARETINA ARVO", SOURCES["FARETINA ARVO"]), ("Gong 2024", SOURCES["Gong 2024"]), ("Gallivan 2023", SOURCES["Gallivan 2023"]), ("Fevrier 2024", SOURCES["Fevrier 2024"])],
    )

    add_section_heading(doc, "Evidence Access and Caveats", 2)
    add_cited_para(
        doc,
        "Open full text or public PDFs were available for several high-value methods sources, including Wykoff 2024, Gallivan 2023, Fevrier 2024, Khanani 2022, the ASRS conversion abstract PDF, the Verana ARVO nAMD+GA presentation, and Retina Society posters.",
        [("Wykoff 2024", SOURCES["Wykoff 2024"]), ("Gallivan 2023", SOURCES["Gallivan 2023"]), ("Fevrier 2024", SOURCES["Fevrier 2024"]), ("Khanani 2022", SOURCES["Khanani 2022"]), ("ASRS 2022", SOURCES["ASRS 2022"]), ("Leng 2024", SOURCES["Leng 2024"]), ("Khurana LTFU", SOURCES["Khurana LTFU"]), ("Saved manifest", SOURCES["Saved manifest"])],
    )
    add_cited_para(
        doc,
        "For some manuscript records, only abstracts or publisher metadata were publicly available in this environment, so those rows are marked accordingly and should be treated as lower-resolution for diagnostic/procedure code detail.",
        [("MacCumber CJO", SOURCES["MacCumber 2023 CJO"]), ("Acharya 2025", SOURCES["Acharya 2025"])],
    )
    add_cited_para(
        doc,
        "Long code lists are not reproduced in full when they reside in article supplements or appendices; instead, this report cites the article table/supplement URL and identifies the local extraction where the pointer was found.",
        [("Gallivan Table A", SOURCES["Gallivan 2023"]), ("Khanani supplement", SOURCES["Khanani supplement"]), ("Wykoff supplement", SOURCES["Wykoff supplement"])],
    )

    # Landscape inventory table
    new_section(doc, True)
    add_section_heading(doc, "Publication Inventory", 3)
    inventory_rows = []
    for item in sorted(manifest, key=lambda x: (x.get("scope", ""), x.get("year", 0), x.get("title", ""))):
        url = best_url(item)
        inventory_rows.append({
            "Year": item.get("year", ""),
            "Scope": capitalize_first(item.get("scope", "")),
            "Item": item.get("title", ""),
            "Format": title_style_label(item.get("format", "")),
            "Qualification": item.get("qualification", ""),
            "Source": (source_label(url), url),
        })
    add_table(doc, ["Year", "Scope", "Item", "Format", "Qualification", "Source"], inventory_rows)
    add_table_footer(
        doc,
        "Scope definitions: Core = publications whose main cohort, exposure, or outcome evaluates patients with wet AMD/nAMD/exudative AMD or anti-VEGF treatment for that condition. Adjacent = IRIS Registry AMD or retinal-injection studies that inform wet AMD evaluation but primarily address conversion risk, geographic atrophy/fellow-eye status, broader AMD categories, or related injection safety rather than a primary wet AMD treatment/outcomes cohort.",
        [("Saved manifest", SOURCES["Saved manifest"]), ("AAO IRIS", AAO)],
    )

    # Detailed operational table
    new_section(doc, True)
    add_section_heading(doc, "Operational Methods Matrix", 4)
    method_rows = [
        {
            "Study": "Wykoff 2024 / Leng 2021 long-term anti-VEGF outcomes",
            "Population / Index": "nAMD treated July 1, 2013-Jun 30, 2018; age >=50; first nAMD diagnosis within 180 days before/on first anti-VEGF injection; index = first documented anti-VEGF injection.",
            "Outcome Definitions": "VA annually +/-60 days; Snellen converted to ETDRS = 85 + 50 x log(Snellen fraction); gap >18 to <=52 weeks; discontinuation >52 weeks; switch = >=3 consecutive injections of another agent; CVL >=10-letter loss; SPV 20/200 or worse twice >=3 months apart without later improvement beyond 20/100.",
            "Analysis": "Descriptive treatment patterns; adjusted linear regression for year-1 VA change; Kaplan-Meier and Cox models for CVL/SPV.",
            "Code / Appendix Pointer": "nAMD ICD-9/10 Table S1 is available in the Wykoff supplementary tables PDF.",
            "Source": [("Wykoff 2024", SOURCES["Wykoff 2024"]), ("Leng 2021 ASRS", SOURCES["Leng 2021 ASRS"]), ("Wykoff supplement", SOURCES["Wykoff supplement"])],
        },
        {
            "Study": "Gallivan 2023 VIEW 1/2 emulation",
            "Population / Index": "Aflibercept/ranibizumab injections Jan 1, 2013-Dec 31, 2018; index = first injection; nAMD diagnosis in same eye within 6 months; baseline VA 20/40-20/320 within 30 days before index.",
            "Outcome Definitions": "Injection = CPT 67028 plus same-day HCPCS J0178/J2778/J3490/J3590; 1-year VA day 365 +/-14 days; maintained vision = <15-letter or <3-line loss.",
            "Analysis": "Complete-case, multiple imputation, and transportability-weighted imputation; VIEW-like treatment regimen assignment.",
            "Code / Appendix Pointer": "Table A includes nAMD ICD-9 362.52 and ICD-10 H35.32/H35.3210-H35.3293 variants plus exclusion codes; local extracted text has full table excerpt.",
            "Source": ("Gallivan 2023", SOURCES["Gallivan 2023"]),
        },
        {
            "Study": "Fevrier 2024 CATT emulation",
            "Population / Index": "Treatment-naive nAMD eyes with first anti-VEGF Oct 1, 2015-Dec 31, 2019; no structured anti-VEGF CPT/HCPCS evidence in prior year; index = first injection.",
            "Outcome Definitions": "PRN = retina-specialist encounter every 4-6 weeks for 1 year plus >=1 noninjection retina encounter; only bevacizumab or ranibizumab for 1 year; VA converted from Snellen using 85 + 50 x log(Snellen fraction).",
            "Analysis": "Exact matching 1:1 to CATT on age, gender, baseline VA; age within 5 years except >=90; VA within 5 letters; IPSW also applied.",
            "Code / Appendix Pointer": "Article states ICD-10/CPT/HCPCS coding but does not print exact code list in extracted text.",
            "Source": ("Fevrier 2024", SOURCES["Fevrier 2024"]),
        },
        {
            "Study": "Khurana 2020/2023 LTFU and nonpersistence",
            "Population / Index": "Treatment-naive nAMD patients treated with anti-VEGF from 2013-2015 and followed through 2019 in manuscript; poster cohort diagnosed 2013-2015 and treated through 2018.",
            "Outcome Definitions": "Manuscript LTFU = no follow-up within 12 months from last intravitreal injection; nonpersistence = no follow-up within 6 months. Poster excluded PDR, DME, RVO, myopic degeneration, idiopathic CNV.",
            "Analysis": "Multivariable logistic regression; covariates included demographics, eye involvement, baseline vision, region, insurance, and provider ZIP-code income in poster.",
            "Code / Appendix Pointer": "No formal public code table found in accessible poster/abstract.",
            "Source": [("Khurana LTFU", SOURCES["Khurana LTFU"]), ("Khurana 2023", SOURCES["Khurana 2023"])],
        },
        {
            "Study": "Khanani 2022 brolucizumab safety",
            "Population / Index": "Adult nAMD patients initiating brolucizumab; IRIS index period Oct 8, 2019-Jun 5, 2020; index = earliest brolucizumab injection by procedure code or EHR note.",
            "Outcome Definitions": "36-month pre-index and <=180-day post-index; IOI and/or RO and RV and/or RO identified by ICD-10-CM; infectious IOI/endophthalmitis excluded; incident logic considered pre-index event history and post-event VA drop >=3 lines.",
            "Analysis": "Patient-eye outcomes; GEE multivariable models for inter-eye correlation; age/sex plus clinically relevant/significant variables.",
            "Code / Appendix Pointer": "eTable 1 inclusion/exclusion; eTable 2 adverse-event definitions; supplement URL cited.",
            "Source": [("Khanani 2022", SOURCES["Khanani 2022"]), ("Khanani supplement", SOURCES["Khanani supplement"])],
        },
        {
            "Study": "Zarbin 2024 brolucizumab safety",
            "Population / Index": "18,312 eyes / 15,998 patients with >=1 brolucizumab injection Oct 8, 2019-Oct 7, 2021; index = first brolucizumab injection.",
            "Outcome Definitions": "AE count/percent, time to AE from index, number of prior brolucizumab injections before AE, VA at/immediately after AE and 6 months after AE.",
            "Analysis": "Patient-eye outcomes; relative risk over time; Cox models adjusted for clustering when both eyes from same patient.",
            "Code / Appendix Pointer": "Supplementary ESM Table S1 is available from the Springer article page.",
            "Source": [("Zarbin 2024", SOURCES["Zarbin 2024"]), ("Zarbin ESM", SOURCES["Zarbin ESM"])],
        },
        {
            "Study": "FARETINA-AMD faricimab outcomes",
            "Population / Index": "nAMD patients initiating faricimab; ARVO analysis used Feb-Aug 2022 starts; later manuscripts extend to 2023.",
            "Outcome Definitions": "Faricimab captured by rules-based EHR text search with regular-expression keywords; required >=12 months pre-initiation EHR data and known laterality; interval/BDVA analyses required >=4 injections; extended interval = any interval >6 weeks.",
            "Analysis": "Retrospective treatment-naive and previously treated cohorts; VA/BDVA, CST, CST <=280 micrometers, injection frequency.",
            "Code / Appendix Pointer": "No public keyword/NDC table found in accessible abstracts.",
            "Source": [("FARETINA ARVO", SOURCES["FARETINA ARVO"]), ("Ali 2025", SOURCES["Ali 2025"]), ("Tabano 2026", SOURCES["Tabano 2026"])],
        },
        {
            "Study": "Gong/ASRS conversion to exudative AMD",
            "Population / Index": "IRIS 2016-2019; 2,664,789 patients with dry AMD in at least one eye; manuscript focuses on unilateral exudative AMD fellow-eye conversion.",
            "Outcome Definitions": "Conversion time from dry/nonexudative to wet/exudative AMD; covariates included age, sex, race, region, smoking, dry AMD stage, wet AMD stage, fellow-eye status.",
            "Analysis": "Descriptive statistics and Cox proportional hazards; manuscript also reports logistic regression.",
            "Code / Appendix Pointer": "No diagnostic-code table found in accessible poster/abstract.",
            "Source": [("Gong 2024", SOURCES["Gong 2024"]), ("ASRS 2022", SOURCES["ASRS 2022"])],
        },
        {
            "Study": "Leng 2024 nAMD + GA anti-VEGF presentation",
            "Population / Index": "First GA ICD-10 code July 1, 2016-Dec 31, 2021; same-eye nAMD ICD-10; cohorts: GA after nAMD, GA before nAMD, coincident GA+nAMD.",
            "Outcome Definitions": "Baseline VA nearest index within 6 months pre-index; treatment year = 52 +/-8 weeks; nearest VA selected, later if equidistant, best same-day VA used; AE within 120 days of injection with no prior history; IOP elevation >6 mmHg and concurrent >=25 mmHg.",
            "Analysis": "Descriptive cohort comparisons by diagnosis timing, anti-VEGF intervals, discontinuation, and safety-event incidence.",
            "Code / Appendix Pointer": "Presentation cites ICD-10/ICD-10-CM but does not list exact codes.",
            "Source": ("Leng 2024", SOURCES["Leng 2024"]),
        },
    ]
    add_table(doc, ["Study", "Population / Index", "Outcome Definitions", "Analysis", "Code / Appendix Pointer", "Source"], method_rows)

    # Recommended design pattern table
    new_section(doc, True)
    add_section_heading(doc, "Recommended Operational Design Pattern", 5)
    design_rows = [
        {
            "Component": "Disease cohort",
            "Recommended operationalization": "Use laterality-specific ICD-coded nAMD/wet AMD diagnosis and require same-eye anti-VEGF evidence when evaluating treated patients.",
            "Why it is preferred": "The largest and most methodologically explicit IRIS studies anchor diagnosis to same-eye treatment and exclude unknown laterality.",
            "Source": ("Wykoff 2024", SOURCES["Wykoff 2024"]),
        },
        {
            "Component": "Index date",
            "Recommended operationalization": "Use first qualifying anti-VEGF injection for treatment-naive outcomes, first drug-specific injection for switch/drug studies, or dry AMD cohort entry for conversion studies.",
            "Why it is preferred": "This reduces ambiguity between prevalent disease, treatment initiation, and drug-switch questions.",
            "Source": [("Gallivan 2023", SOURCES["Gallivan 2023"]), ("Khanani 2022", SOURCES["Khanani 2022"]), ("MacCumber interval", SOURCES["MacCumber 2023 interval"]), ("Gong 2024", SOURCES["Gong 2024"]), ("ASRS 2022", SOURCES["ASRS 2022"]), ("Research Square", SOURCES["Research Square"])],
        },
        {
            "Component": "Lookback / treatment naive",
            "Recommended operationalization": "Require pre-index data contribution and no structured anti-VEGF evidence in the lookback when claiming treatment-naive status.",
            "Why it is preferred": "Fevrier used 1-year pre-index anti-VEGF-free evidence; Wykoff required 6 months of practice contribution before index.",
            "Source": [("Fevrier 2024", SOURCES["Fevrier 2024"]), ("Wykoff 2024", SOURCES["Wykoff 2024"])],
        },
        {
            "Component": "Visual acuity",
            "Recommended operationalization": "Predefine baseline/follow-up windows and convert Snellen to ETDRS-equivalent letters using published formula; choose best/nearest measurement rules up front.",
            "Why it is preferred": "This is the most common VA handling pattern and allows comparability across long-term outcomes and trial-emulation analyses.",
            "Source": [("Wykoff 2024", SOURCES["Wykoff 2024"]), ("Leng 2024", SOURCES["Leng 2024"])],
        },
        {
            "Component": "Treatment burden",
            "Recommended operationalization": "Capture injection frequency, injection interval, gaps, discontinuation, switching, and persistence/LTFU.",
            "Why it is preferred": "These endpoints repeatedly explain real-world under-treatment and outcome differences.",
            "Source": [("Wykoff 2024", SOURCES["Wykoff 2024"]), ("MacCumber CJO", SOURCES["MacCumber 2023 CJO"]), ("Khurana 2023", SOURCES["Khurana 2023"])],
        },
        {
            "Component": "Statistical model",
            "Recommended operationalization": "Use linear/logistic/Poisson models for cross-sectional or 1-year outcomes, Cox models for time-to-event outcomes, and account for two-eye correlation.",
            "Why it is preferred": "GEE or clustered Cox approaches address nonindependence when both eyes contribute data.",
            "Source": [("Khanani 2022", SOURCES["Khanani 2022"]), ("Gallivan 2023", SOURCES["Gallivan 2023"]), ("Zarbin 2024", SOURCES["Zarbin 2024"]), ("Wykoff 2024", SOURCES["Wykoff 2024"]), ("Acharya 2025", SOURCES["Acharya 2025"])],
        },
        {
            "Component": "Codes and appendices",
            "Recommended operationalization": "Put full ICD/CPT/HCPCS/NDC or keyword lists in an appendix and cite them from the methods table.",
            "Why it is preferred": "Code tables can be too long for the main report; Gallivan Table A and Khanani eTables are good precedents.",
            "Source": [("Gallivan 2023", SOURCES["Gallivan 2023"]), ("Khanani supplement", SOURCES["Khanani supplement"])],
        },
    ]
    add_table(doc, ["Component", "Recommended operationalization", "Why it is preferred", "Source"], design_rows)

    # Portrait conclusion/source pointers
    new_section(doc, False)
    add_section_heading(doc, "Implications for New IRIS Wet Age-related Macular Degeneration Studies", 6)
    bullet_num_id = add_bullet_numbering(doc)
    add_cited_para(
        doc,
        "A new wet AMD study should begin with a precise estimand: treatment initiation, treatment burden, treatment response, switch durability, safety, conversion, or disparities. Each estimand maps to a different index date and follow-up window.",
        [("Wykoff 2024", SOURCES["Wykoff 2024"]), ("MacCumber interval", SOURCES["MacCumber 2023 interval"]), ("Khanani 2022", SOURCES["Khanani 2022"]), ("Gong 2024", SOURCES["Gong 2024"]), ("Acharya 2025", SOURCES["Acharya 2025"]), ("Gallivan 2023", SOURCES["Gallivan 2023"]), ("Fevrier 2024", SOURCES["Fevrier 2024"])],
        bullet_num_id=bullet_num_id,
    )
    add_cited_para(
        doc,
        "For treatment-response questions, the strongest recurring approach is an eye-level treated nAMD cohort with baseline VA, minimum exposure/follow-up, injection-count endpoints, and adjustment for age, baseline VA, payer, race/ethnicity, provider specialty, and ocular comorbidity.",
        [("Wykoff 2024", SOURCES["Wykoff 2024"]), ("Barikian 2026", SOURCES["Barikian 2026"])],
        bullet_num_id=bullet_num_id,
    )
    add_cited_para(
        doc,
        "For adherence questions, LTFU and nonpersistence should be separated because the reviewed manuscript used 12-month no-follow-up for LTFU and 6-month no-follow-up for nonpersistence.",
        [("Khurana 2023", SOURCES["Khurana 2023"])],
        bullet_num_id=bullet_num_id,
    )
    add_cited_para(
        doc,
        "For safety questions, code-based adverse events should be presented with enough specificity to distinguish infectious events, IOI, retinal vasculitis, retinal vascular occlusion, and post-event VA change.",
        [("Khanani 2022", SOURCES["Khanani 2022"]), ("Zarbin 2024", SOURCES["Zarbin 2024"])],
        bullet_num_id=bullet_num_id,
    )

    new_section(doc, True)
    add_section_heading(doc, "Citation Register", 7)
    source_rows = [
        {"Citation": CITATION_TEXTS.get(label, label), "Use in report": "Primary linked source for cited statements/tables.", "URL": ("open", url)}
        for label, url in SOURCES.items()
    ]
    add_table(doc, ["Citation", "Use in report", "URL"], source_rows)

    # Save
    clear_headers_and_footers(doc)
    doc.core_properties.title = "Evaluating Patients with Wet Age-related Macular Degeneration Using the IRIS Registry"
    doc.core_properties.subject = "Methods review with inline citations"
    doc.core_properties.author = "OpenAI Codex"
    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
