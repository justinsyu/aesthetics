from __future__ import annotations

import csv
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
TEMPLATE = Path(r"C:\Users\Justin\Desktop\LinkedIn\Target_Product_Profile_cohere_report.docx")
OUTPUT = BASE / "Vestrum_wet_AMD_methods_comparison_report.docx"
CSV_PATH = BASE / "methods_comparison.csv"

DARK = "10120F"
BODY = "1B1F17"
MUTED = "5C6257"
ACCENT = "B7E76B"
HEADER_TEXT = "D7FF5F"
CREAM_TEXT = "F6F1E8"
LIGHT = "F6F1E8"
TABLE_ALT = "EBE4D6"
BORDER = "CBD2C2"
LINK = "2A5DB0"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_document_background(doc: Document, fill: str = CREAM_TEXT) -> None:
    root = doc._element
    background = root.find(qn("w:background"))
    if background is None:
        background = OxmlElement("w:background")
        root.insert(0, background)
    background.set(qn("w:color"), fill)


def set_run_char_spacing(run, value: int) -> None:
    rpr = run._r.get_or_add_rPr()
    spacing = rpr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        rpr.append(spacing)
    spacing.set(qn("w:val"), str(value))


def set_run_shading(run, fill: str) -> None:
    rpr = run._r.get_or_add_rPr()
    shd = rpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        rpr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_keep_next_and_bottom_rule(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    if ppr.find(qn("w:keepNext")) is None:
        ppr.append(OxmlElement("w:keepNext"))
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = pbdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pbdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), DARK)


def set_margins(section) -> None:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    sect_pr = section._sectPr
    shd = sect_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        sect_pr.append(shd)
    shd.set(qn("w:fill"), CREAM_TEXT)


def set_orientation(section, orientation: str) -> None:
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    set_margins(section)


def replace_paragraph_text_preserve_first_run(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def update_headers_and_footers(doc: Document) -> None:
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip():
                replace_paragraph_text_preserve_first_run(paragraph, "Vestrum Wet AMD Methods Comparison")


def add_hyperlink(paragraph, text: str, url: str, bold: bool = False, italic: bool = False, color_value: str = LINK) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_value)
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    if bold:
        rpr.append(OxmlElement("w:b"))
    if italic:
        rpr.append(OxmlElement("w:i"))

    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def format_para(paragraph, before: float = 0, after: float = 6, line: float = 1.05) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade_paragraph(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_label(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    format_para(p, after=8, line=1.25)
    run = p.add_run(text)
    run.bold = True
    run.font.all_caps = True
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(DARK)
    set_run_char_spacing(run, 30)
    set_run_shading(run, HEADER_TEXT)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    format_para(p, after=7, line=1.3333333333333333)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = rgb(DARK)


def add_subtitle(doc: Document, text: str, citations: list[tuple[str, str]] | None = None) -> None:
    p = doc.add_paragraph()
    format_para(p, after=10, line=1.3333333333333333)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = rgb(MUTED)
    add_citations(p, citations or [])


def add_heading(doc: Document, text: str, n: int | None = None) -> None:
    p = doc.add_paragraph()
    format_para(p, before=18, after=8, line=1.3333333333333333)
    set_keep_next_and_bottom_rule(p)
    label = f"{n}.  {text}" if n is not None else text
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = rgb(DARK)


def add_citations(paragraph, citations: list[tuple[str, str]], color_value: str = LINK) -> None:
    if not citations:
        return
    paragraph.add_run(" ")
    paragraph.add_run("(")
    for i, (label, url) in enumerate(citations):
        if i:
            paragraph.add_run("; ")
        add_hyperlink(paragraph, label, url, color_value=color_value)
    paragraph.add_run(")")


def add_body_paragraph(doc: Document, text: str, citations: list[tuple[str, str]]) -> None:
    p = doc.add_paragraph()
    format_para(p, after=10, line=1.25)
    run = p.add_run(text)
    run.font.color.rgb = rgb(BODY)
    add_citations(p, citations)


def add_bullet(doc: Document, text: str, citations: list[tuple[str, str]], callout: bool = False) -> None:
    p = doc.add_paragraph()
    format_para(p, after=8, line=1.3333333333333333)
    p.paragraph_format.left_indent = Pt(23)
    p.paragraph_format.first_line_indent = Pt(-14)
    p.paragraph_format.tab_stops.add_tab_stop(Pt(23), WD_TAB_ALIGNMENT.LEFT)
    bullet = p.add_run("\u2014\t")
    bullet.font.color.rgb = rgb(CREAM_TEXT if callout else BODY)
    run = p.add_run(text)
    run.font.color.rgb = rgb(CREAM_TEXT if callout else BODY)
    if callout:
        shade_paragraph(p, DARK)
    add_citations(p, citations, color_value=HEADER_TEXT if callout else LINK)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = BODY, outer_size: str = "4") -> None:
    # The reference DOCX relies on table-level borders rather than per-cell borders.
    return
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), outer_size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 90, bottom: int = 90, left: int = 110, right: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    look = tbl_pr.find(qn("w:tblLook"))
    if look is None:
        look = OxmlElement("w:tblLook")
        tbl_pr.append(look)
    look.set(qn("w:val"), "0000")
    look.set(qn("w:firstRow"), "0")
    look.set(qn("w:lastRow"), "0")
    look.set(qn("w:firstColumn"), "0")
    look.set(qn("w:lastColumn"), "0")
    look.set(qn("w:noHBand"), "0")
    look.set(qn("w:noVBand"), "0")
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge, size in (("top", "4"), ("left", "4"), ("bottom", "4"), ("right", "4"), ("insideH", "2"), ("insideV", "2")):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BODY)


def set_cell_text(
    cell,
    text: str,
    bold: bool = False,
    size: float = 7.4,
    fill: str | None = None,
    color: str = DARK,
    caps: bool = False,
    char_spacing: int | None = None,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    format_para(p, after=0, line=1.12)
    r = p.add_run(text)
    r.bold = bold
    r.font.all_caps = caps
    r.font.size = Pt(size)
    r.font.color.rgb = rgb(color)
    if char_spacing is not None:
        set_run_char_spacing(r, char_spacing)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_borders(cell)
    set_cell_margins(cell)
    if fill:
        set_cell_shading(cell, fill)


def set_source_cell(cell, citations: list[tuple[str, str]], size: float = 7.4) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    format_para(p, after=0, line=1.0)
    for i, (label, url) in enumerate(citations):
        if i:
            p.add_run().add_break()
        add_hyperlink(p, label, url)
    for run in p.runs:
        run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_borders(cell)
    set_cell_margins(cell)


def shade_row(row, fill: str) -> None:
    for cell in row.cells:
        set_cell_shading(cell, fill)


def source_label(url: str) -> str:
    host = re.sub(r"^https?://(www\.)?", "", url).split("/")[0].lower()
    if "pubmed" in url:
        return "PubMed"
    if "pmc.ncbi" in url:
        return "PMC"
    if "sciencedirect" in url:
        return "ScienceDirect"
    if "sagepub" in url:
        return "Sage"
    if "gene.com" in url:
        return "Poster/PDF"
    if "retinalphysician" in url:
        return "Retinal Physician"
    if "nature.com" in url:
        return "Nature"
    if "vestrumhealth" in url:
        return "Vestrum"
    return host


def row_citations(row: dict, max_links: int = 2) -> list[tuple[str, str]]:
    urls = [u.strip() for u in row.get("source_urls", "").split(";") if u.strip()]
    cites = []
    for url in urls[:max_links]:
        cites.append((f"{row['year']} {source_label(url)}", url))
    return cites


def add_landscape_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_orientation(section, "landscape")


def add_portrait_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_orientation(section, "portrait")


def add_table_heading(doc: Document, text: str, citations: list[tuple[str, str]] | None = None) -> None:
    p = doc.add_paragraph()
    format_para(p, before=18, after=8, line=1.3333333333333333)
    set_keep_next_and_bottom_rule(p)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = rgb(DARK)
    add_citations(p, citations or [])


def build_question_table(doc: Document, rows: list[dict]) -> None:
    add_landscape_section(doc)
    add_table_heading(doc, "Method Selection Matrix", [
        ("Ciulla 2020", "https://pubmed.ncbi.nlm.nih.gov/31324588/"),
        ("SIERRA-AMD", "https://pubmed.ncbi.nlm.nih.gov/31812631/"),
        ("Ko 2026", "https://medically.gene.com/content/dam/pdmahub/restricted/ophthalmology/amcp-2026/AMCP-2026-poster-ko-real-world-persistence-switching-and-reinitiation-in-patients.pdf"),
    ])
    data = [
        ("First-year VA or treatment intensity", "Treatment-naive nAMD eyes; first injection as index; baseline VA strata; injection-count exposure; month-12/1-year VA endpoint.", "Ciulla 2020; Moshfeghi 2021; Ciulla 2018", [("Ciulla 2020", "https://pubmed.ncbi.nlm.nih.gov/31324588/"), ("Moshfeghi 2021", "https://pmc.ncbi.nlm.nih.gov/articles/PMC9979036/")]),
        ("Long-term treatment burden", "First injection index; annual/cumulative injection counts; visits and noninjection visits; years 1-4 or 1-5 outcomes where follow-up exists.", "SIERRA-AMD; Ciulla 2022", [("SIERRA-AMD", "https://pubmed.ncbi.nlm.nih.gov/31812631/"), ("Ciulla 2022", "https://pubmed.ncbi.nlm.nih.gov/35381391/")]),
        ("Persistence, gaps, switching, or reinitiation", "Use a clear pre-index history window and at least 24 months follow-up when feasible; define nonpersistence/gap as >=180 days.", "Ko 2026; Moshfeghi 2026", [("Ko 2026", "https://medically.gene.com/content/dam/pdmahub/restricted/ophthalmology/amcp-2026/AMCP-2026-poster-ko-real-world-persistence-switching-and-reinitiation-in-patients.pdf"), ("Moshfeghi 2026", "https://www.retinalphysician.com/issues/2026/may-june/arvo32/")]),
        ("Agent comparison or modern durability", "Index on first use for treatment-naive eyes; use switch date for treatment-experienced eyes; adjust or stratify by baseline/switch-date VA; report interval between final injections.", "Rowe 2026; Ko 2026", [("Rowe 2026", "https://journals.sagepub.com/doi/10.1177/24741264261428749"), ("Ko 2026", "https://medically.gene.com/content/dam/pdmahub/restricted/ophthalmology/amcp-2026/AMCP-2026-poster-ko-real-world-persistence-switching-and-reinitiation-in-patients.pdf")]),
        ("Conversion or prevention", "For dry-to-wet conversion, require diagnostic-code change plus anti-VEGF initiation when feasible; use earliest event date as conversion date.", "Luttrull 2023; Starr 2021", [("Luttrull 2023", "https://pmc.ncbi.nlm.nih.gov/articles/PMC10550910/"), ("Starr 2021", "https://pubmed.ncbi.nlm.nih.gov/34038686/")]),
        ("Functional or adverse-event endpoints", "Use endpoint-specific rules: driving vision loss as VA worse than 20/40 sustained >=6 months; SMH rate by injection type and timing versus prior injection.", "Emami 2024; Kaufmann 2025", [("Emami 2024", "https://pubmed.ncbi.nlm.nih.gov/37866681/"), ("Kaufmann 2025", "https://pubmed.ncbi.nlm.nih.gov/39455036/")]),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 13680)
    widths = [2500, 5600, 2600, 2980]
    headers = ["Research question", "Recommended operational design", "Closest Vestrum templates", "Inline sources"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=8.0, fill=DARK, color=HEADER_TEXT, caps=True, char_spacing=12)
        set_cell_width(table.rows[0].cells[i], widths[i])
    for row_idx, (question, design, templates, citations) in enumerate(data, start=1):
        row = table.add_row()
        cells = row.cells
        shade_row(row, LIGHT if row_idx % 2 else TABLE_ALT)
        for i, cell in enumerate(cells):
            set_cell_width(cell, widths[i])
        set_cell_text(cells[0], question, bold=True, size=8.0)
        set_cell_text(cells[1], design, size=8.0)
        set_cell_text(cells[2], templates, size=8.0)
        set_source_cell(cells[3], citations)


def build_publication_table(doc: Document, rows: list[dict]) -> None:
    add_landscape_section(doc)
    add_table_heading(doc, "Publication-Level Methods Comparison", [
        ("Vestrum Media", "https://www.vestrumhealth.com/media.php"),
        ("Detailed extraction", "file:///" + str((BASE / "methods_comparison.md").resolve()).replace("\\", "/")),
    ])
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 13680)
    widths = [2150, 2800, 2500, 2800, 2850, 1380]
    headers = ["Publication", "Population / index", "Exposure / follow-up", "Operational outcomes", "Code / supplement status", "Sources"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=6.5, fill=DARK, color=HEADER_TEXT, caps=True, char_spacing=4)
        set_cell_width(table.rows[0].cells[i], widths[i])
    for row_idx, row in enumerate(rows, start=1):
        table_row = table.add_row()
        cells = table_row.cells
        shade_row(table_row, LIGHT if row_idx % 2 else TABLE_ALT)
        for i, cell in enumerate(cells):
            set_cell_width(cell, widths[i])
        title = re.sub(r"^\d+\.\s*", "", row["title"])
        set_cell_text(cells[0], f"{title} ({row['year']})", bold=True, size=6.5)
        set_cell_text(cells[1], f"{row['population_inclusion']} {row['index_exposure_definition']}", size=6.7)
        set_cell_text(cells[2], f"{row['follow_up_windows']} {row['exclusions']}", size=6.7)
        set_cell_text(cells[3], row["outcome_operational_definitions"], size=6.7)
        set_cell_text(cells[4], f"{row['diagnosis_case_definition_and_codes']} {row['treatment_drug_definition']} {row['code_lists_appendices_or_supplements']}", size=6.7)
        set_source_cell(cells[5], row_citations(row), size=6.7)


def build_source_list(doc: Document, rows: list[dict]) -> None:
    add_portrait_section(doc)
    add_heading(doc, "Source Links", 6)
    add_body_paragraph(doc, "The links below are the publication landing pages, abstracts, full-text records, posters, or source pages used for inline citation in this report.", [
        ("Vestrum Media", "https://www.vestrumhealth.com/media.php"),
        ("Methods extraction file", "file:///" + str((BASE / "methods_comparison.md").resolve()).replace("\\", "/")),
    ])
    seen = set()
    for row in rows:
        title = re.sub(r"^\d+\.\s*", "", row["title"])
        urls = [u.strip() for u in row.get("source_urls", "").split(";") if u.strip()]
        if not urls:
            continue
        p = doc.add_paragraph()
        format_para(p, after=6, line=1.25)
        p.paragraph_format.left_indent = Pt(23)
        p.paragraph_format.first_line_indent = Pt(-14)
        p.paragraph_format.tab_stops.add_tab_stop(Pt(23), WD_TAB_ALIGNMENT.LEFT)
        b = p.add_run("\u2014\t")
        b.font.size = Pt(8.5)
        b.font.color.rgb = rgb(DARK)
        r = p.add_run(f"{title} ({row['year']}): ")
        r.font.size = Pt(8.5)
        r.font.color.rgb = rgb(DARK)
        for i, url in enumerate(urls):
            if i:
                p.add_run("; ")
            label = source_label(url)
            add_hyperlink(p, label, url)
            seen.add(url)


def main() -> None:
    rows = list(csv.DictReader(CSV_PATH.open(encoding="utf-8")))
    shutil.copyfile(TEMPLATE, OUTPUT)
    doc = Document(OUTPUT)
    clear_document(doc)
    set_document_background(doc)
    set_orientation(doc.sections[0], "portrait")

    add_label(doc, "Vestrum Health  \u00b7  Wet AMD Methods Evidence Review")
    add_title(doc, "Vestrum Wet AMD Methods Comparison")
    add_subtitle(doc, "A source-linked synthesis of Vestrum database studies evaluating neovascular/wet AMD, focused on reusable cohort definitions, treatment exposure rules, outcomes, and reproducibility gaps.", [
        ("Vestrum Media", "https://www.vestrumhealth.com/media.php"),
        ("Methods extraction", "file:///" + str((BASE / "methods_comparison.md").resolve()).replace("\\", "/")),
    ])

    add_heading(doc, "Executive Takeaways", 1)
    add_bullet(doc, "The most reusable default design is an eye-level retrospective Vestrum cohort using nAMD diagnosis plus anti-VEGF treatment, first anti-VEGF injection as index, baseline VA availability, a defined treatment-history window, fixed follow-up/visit windows, and explicit attrition handling.", [
        ("Ciulla 2018", "https://pubmed.ncbi.nlm.nih.gov/31047372/"),
        ("Ciulla 2020", "https://pubmed.ncbi.nlm.nih.gov/31324588/"),
        ("Ciulla 2022", "https://pubmed.ncbi.nlm.nih.gov/35381391/"),
    ], callout=True)
    add_bullet(doc, "Treatment exposure is most commonly operationalized with annual/cumulative injection counts, inter-injection intervals, and treatment gaps; recent persistence/gap analyses converge on a >=180-day gap threshold.", [
        ("SIERRA-AMD", "https://pubmed.ncbi.nlm.nih.gov/31812631/"),
        ("Ko 2026", "https://medically.gene.com/content/dam/pdmahub/restricted/ophthalmology/amcp-2026/AMCP-2026-poster-ko-real-world-persistence-switching-and-reinitiation-in-patients.pdf"),
        ("Moshfeghi 2026", "https://www.retinalphysician.com/issues/2026/may-june/arvo32/"),
    ], callout=True)
    add_bullet(doc, "VA outcomes should be reported as ETDRS-equivalent letters with a stated conversion method and baseline VA strata because baseline VA strongly affects apparent gains and losses.", [
        ("Ciulla 2018", "https://pubmed.ncbi.nlm.nih.gov/31047372/"),
        ("Moshfeghi 2021", "https://pmc.ncbi.nlm.nih.gov/articles/PMC9979036/"),
        ("Ciulla 2022", "https://pubmed.ncbi.nlm.nih.gov/35381391/"),
    ], callout=True)
    add_bullet(doc, "The most explicit published conversion algorithm requires both diagnostic-code change and anti-VEGF initiation, with conversion dated to the earliest of those events; this is strongest for dry-to-wet conversion work.", [
        ("Luttrull 2023", "https://pmc.ncbi.nlm.nih.gov/articles/PMC10550910/"),
        ("Starr 2021", "https://pubmed.ncbi.nlm.nih.gov/34038686/"),
    ], callout=True)
    add_bullet(doc, "The major reproducibility gap is that most accessible publications do not publish complete ICD/NDC/HCPCS lists, so any future Vestrum analysis should preserve those lists internally and publish a supplement or appendix pointer when feasible.", [
        ("Moshfeghi 2021", "https://pmc.ncbi.nlm.nih.gov/articles/PMC9979036/"),
        ("Luttrull 2023", "https://pmc.ncbi.nlm.nih.gov/articles/PMC10550910/"),
        ("Rowe 2026", "https://journals.sagepub.com/doi/10.1177/24741264261428749"),
    ], callout=True)

    add_heading(doc, "Recommended Default Design", 2)
    add_body_paragraph(doc, "For a treatment-outcomes analysis, start with treatment-naive nAMD eyes, define first anti-VEGF injection as the index date, require baseline VA, use a defined pre-index history window, and specify fixed follow-up windows before analysis. This design is most consistent with the Ciulla first-year and longer-term analyses, the frequent-treatment analysis, and SIERRA-AMD.", [
        ("Ciulla 2020", "https://pubmed.ncbi.nlm.nih.gov/31324588/"),
        ("Moshfeghi 2021", "https://pmc.ncbi.nlm.nih.gov/articles/PMC9979036/"),
        ("SIERRA-AMD", "https://pubmed.ncbi.nlm.nih.gov/31812631/"),
    ])
    add_body_paragraph(doc, "Use baseline VA strata at minimum as 20/40 or better, 20/41-20/70, 20/71-20/200, and 20/201 or worse; report annual injection counts, cumulative injection counts, treatment intervals, and treatment-gap status alongside VA change.", [
        ("Ciulla 2018", "https://pubmed.ncbi.nlm.nih.gov/31047372/"),
        ("Ciulla 2022", "https://pubmed.ncbi.nlm.nih.gov/35381391/"),
        ("SIERRA-AMD", "https://pubmed.ncbi.nlm.nih.gov/31812631/"),
    ])
    add_body_paragraph(doc, "If a functional endpoint is needed, driving vision is the clearest published construct: baseline VA 20/40 or better in the better-seeing eye, with loss defined as VA worse than 20/40 sustained for at least 6 consecutive months.", [
        ("Emami 2024", "https://pubmed.ncbi.nlm.nih.gov/37866681/"),
    ])

    add_heading(doc, "Core Operational Rules", 3)
    add_bullet(doc, "Indexing: first anti-VEGF injection for treated nAMD outcomes; switch date for switcher analyses; first diagnosis for epidemiology or noninjected cohorts.", [
        ("SIERRA-AMD", "https://pubmed.ncbi.nlm.nih.gov/31812631/"),
        ("Rowe 2026", "https://journals.sagepub.com/doi/10.1177/24741264261428749"),
    ])
    add_bullet(doc, "Treatment-naive history: choose one rule and state it; published options include no prior anti-VEGF ever, no anti-VEGF for >180 days before index, or no anti-VEGF during the 12-month pre-index period.", [
        ("SIERRA-AMD", "https://pubmed.ncbi.nlm.nih.gov/31812631/"),
        ("Ko 2026", "https://medically.gene.com/content/dam/pdmahub/restricted/ophthalmology/amcp-2026/AMCP-2026-poster-ko-real-world-persistence-switching-and-reinitiation-in-patients.pdf"),
    ])
    add_bullet(doc, "VA handling: convert Snellen/logMAR to ETDRS-equivalent letters using a stated rule, commonly 85 + 50 x log(Snellen fraction) or 85 - 50 x logMAR; identify whether VA is distance-corrected, near-corrected, pinhole, or mixed real-world VA.", [
        ("Ciulla 2018", "https://pubmed.ncbi.nlm.nih.gov/31047372/"),
        ("Moshfeghi 2021", "https://pmc.ncbi.nlm.nih.gov/articles/PMC9979036/"),
        ("Ciulla 2022", "https://pubmed.ncbi.nlm.nih.gov/35381391/"),
    ])
    add_bullet(doc, "Attrition: define fixed windows and report loss to follow-up separately from observed VA outcomes; the Ciulla 2018 6-, 12-, and 24-month cohorts are the clearest template.", [
        ("Ciulla 2018", "https://pubmed.ncbi.nlm.nih.gov/31047372/"),
    ])

    build_question_table(doc, rows)
    build_publication_table(doc, rows)

    add_portrait_section(doc)
    add_heading(doc, "Evidence Gaps and Caveats", 5)
    add_body_paragraph(doc, "Accessible Vestrum publications generally provide strong outcome-window and treatment-exposure definitions but do not provide complete diagnosis, procedure, and drug-code lists. That limits exact reproducibility even where the analytic design is clear.", [
        ("Detailed extraction", "file:///" + str((BASE / "methods_comparison.md").resolve()).replace("\\", "/")),
        ("Vestrum Media", "https://www.vestrumhealth.com/media.php"),
    ])
    add_body_paragraph(doc, "Several potentially relevant supplements remained blocked or incomplete in this evidence set: the Moshfeghi 2021 supplement endpoint returned a proof-of-work stub, Rowe 2026 Sage supplements returned 403 and PubMed lists PMCID availability for Apr 9 2027, and Kaufmann 2025/AJO supplemental snippets did not expose ICD/HCPCS/NDC lists.", [
        ("Moshfeghi 2021", "https://pmc.ncbi.nlm.nih.gov/articles/PMC9979036/"),
        ("Rowe 2026", "https://journals.sagepub.com/doi/10.1177/24741264261428749"),
        ("Kaufmann 2025", "https://pubmed.ncbi.nlm.nih.gov/39455036/"),
    ])
    add_body_paragraph(doc, "For future Vestrum analyses, the methods supplement should include ICD-9/ICD-10 disease codes, anti-VEGF NDC/HCPCS/procedure codes, exclusion-code lists, VA conversion and visit-window rules, lookback windows, censoring rules, and gap/persistence definitions.", [
        ("Luttrull 2023", "https://pmc.ncbi.nlm.nih.gov/articles/PMC10550910/"),
        ("Ko 2026", "https://medically.gene.com/content/dam/pdmahub/restricted/ophthalmology/amcp-2026/AMCP-2026-poster-ko-real-world-persistence-switching-and-reinitiation-in-patients.pdf"),
    ])

    build_source_list(doc, rows)
    doc.core_properties.title = "Vestrum Wet AMD Methods Comparison"
    doc.core_properties.author = "OpenAI Codex"
    update_headers_and_footers(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
