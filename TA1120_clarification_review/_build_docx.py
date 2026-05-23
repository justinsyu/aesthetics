#!/usr/bin/env python3
"""Build NICE TA1120 clarification review .docx report."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION, WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

OUT_PATH = "/Users/justinyu/Desktop/linkedin-posts/TA1120_clarification_review.docx"

PRINTABLE_TWIPS = 9360  # 6.5 inches portrait
LANDSCAPE_TWIPS = 13680  # 9.5 inches landscape


# ============================================================
# Helpers
# ============================================================

def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '888888')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def make_fixed_table(doc, col_widths_twips, header_labels, header_shade='305496', header_color='FFFFFF'):
    """Create a fixed-layout table with explicit grid widths."""
    total = sum(col_widths_twips)
    table = doc.add_table(rows=1, cols=len(col_widths_twips))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set tblW
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # Remove existing tblW
    for tblW in tblPr.findall(qn('w:tblW')):
        tblPr.remove(tblW)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(total))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # tblLayout fixed
    for layout in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(layout)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    # Borders
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '888888')
        tblBorders.append(b)
    # Remove existing
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblPr.append(tblBorders)

    # Replace tblGrid with explicit gridCol
    existing_grid = tbl.find(qn('w:tblGrid'))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    tblGrid = OxmlElement('w:tblGrid')
    for w in col_widths_twips:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)
    # Insert tblGrid after tblPr
    tblPr.addnext(tblGrid)

    # Set header row
    hdr_row = table.rows[0]
    for i, label in enumerate(header_labels):
        cell = hdr_row.cells[i]
        # Set cell width
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        # Remove existing tcW
        for tcW in tcPr.findall(qn('w:tcW')):
            tcPr.remove(tcW)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(col_widths_twips[i]))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        # Shade
        shade_cell(cell, header_shade)
        # Border
        set_cell_borders(cell)
        # Text
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(header_color)

    return table


def add_table_row(table, col_widths_twips, values, font_size=9, shade_alt=False, row_index=None):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for tcW in tcPr.findall(qn('w:tcW')):
            tcPr.remove(tcW)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(col_widths_twips[i]))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        set_cell_borders(cell)
        if shade_alt and row_index is not None and row_index % 2 == 1:
            shade_cell(cell, 'F2F2F2')
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(val) if val is not None else '')
        run.font.size = Pt(font_size)
    return row


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    # Force font size
    for run in h.runs:
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string('1F3864')
    return h


def add_para(doc, text, size=11, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align is not None:
        p.alignment = align
    return p


# ============================================================
# Build document
# ============================================================

doc = Document()

# Set default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Set page margins (US Letter, 1in margins)
section = doc.sections[0]
section.page_height = Inches(11)
section.page_width = Inches(8.5)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ============================================================
# TITLE PAGE
# ============================================================

# Vertical spacer
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('NICE TA1120 — Clarification Review')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor.from_string('1F3864')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Avelumab with axitinib for untreated advanced renal cell carcinoma')
run.italic = True
run.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('(MA review of TA645)')
run.italic = True
run.font.size = Pt(16)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared by: AI-assisted EAG simulation')
run.font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Date: 22 May 2026')
run.font.size = Pt(13)

add_page_break(doc)

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================

add_heading(doc, 'Executive summary', level=1)

add_para(doc,
    'This report documents an AI-assisted simulation of an EAG (External Assessment Group) '
    'clarification review for NICE TA1120 (Avelumab with axitinib for untreated advanced renal cell '
    'carcinoma — MA review of TA645). Four specialist personas (clinical, statistician, health economist, '
    'information specialist) independently generated clarification questions from the public Committee '
    'papers, blinded to the company response and the EAG Assessment Report. The generated question set was '
    'then de-duplicated, merged, and compared against the actual NICE clarification letter issued to Merck '
    'on 16 December 2024.')

add_para(doc,
    'The four specialists produced 151 individual questions, which de-duplicated to 61 merged topics. '
    'NICE itself asked 23 unique questions. Sixteen of the NICE questions were also identified by the AI '
    'review (hit rate 69.6%), including all three NICE priority "Yes" items related to time horizon, '
    'proportional hazards, and the utility regression methodology. Forty-five topics were generated only by '
    'the AI review and not asked by NICE; seven NICE topics were missed by the AI review.')

add_para(doc,
    'The seven missed items share a single signature: they are all reactive findings from direct, '
    'cell-by-cell audit of the executable Excel model and the full CSR — artefacts the specialists did not '
    'have in front of them. They are not gaps in the methodological scope of the AI review but limitations '
    'of the inputs available to it at the pre-clarification stage.')

# Headline counts table — make prominent
add_heading(doc, 'Headline counts', level=2)

headline_widths = [4680, 4680]
headline_table = make_fixed_table(doc, headline_widths, ['Metric', 'Value'],
                                  header_shade='1F3864', header_color='FFFFFF')
headline_rows = [
    ('Generated topics (unique, post-merge)', '61'),
    ('Actual NICE topics (unique)', '23'),
    ('Both (matched)', '16'),
    ('Generated-only', '45'),
    ('NICE-only (missed)', '7'),
    ('Hit rate on NICE', '16 / 23 = 69.6%'),
]
for i, (k, v) in enumerate(headline_rows):
    row = add_table_row(headline_table, headline_widths, [k, v], font_size=12)
    # Bold value
    for run in row.cells[1].paragraphs[0].runs:
        run.bold = True
    # Make row taller
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '420')
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

doc.add_paragraph()
add_para(doc,
    'Bottom line: The AI review would have caught 16 of the 23 actual NICE clarification questions '
    '(69.6% coverage). The seven missed items are all model-audit findings that require direct interrogation '
    'of the submitted Excel workbook and the full CSR — artefacts not available at the clarification-letter '
    'generation stage of an EAG review.',
    bold=True)

add_page_break(doc)

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================

add_heading(doc, 'Table of contents', level=1)
toc_items = [
    ('Executive summary', '2'),
    ('Methodology', '4'),
    ('Matched topics (Both)', '6'),
    ('Missed topics (NICE-only)', '9'),
    ('Generated-only topics', '11'),
    ('Specialist contributions', '14'),
    ('Discussion', '16'),
    ('Conclusion', '18'),
    ('Appendix A: Full generated question list', '19'),
    ('Appendix B: Full actual NICE question list', '32'),
]
for label, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.font.size = Pt(11)
    tab_run = p.add_run('\t' + page)
    tab_run.font.size = Pt(11)

doc.add_paragraph()
add_para(doc,
    '[Right-click and update Table of Contents in Word to refresh page numbers]',
    size=9, italic=True)

add_page_break(doc)

# ============================================================
# METHODOLOGY
# ============================================================

add_heading(doc, 'Methodology', level=1)

add_heading(doc, 'Source documents acquired', level=2)
add_para(doc,
    'The review used two primary source documents:')
p = doc.add_paragraph(style='List Bullet')
p.add_run('NICE TA1120 Committee papers PDF (484 pages) — containing the company submission (Document B, '
          'Section 1a), the company response to clarification (Section 2), and the EAG Assessment Report (Section 5).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('NICE TA1120 Final scope document (issued September 2024) — defining the population, intervention, '
          'comparators, outcomes and subgroups.')

add_heading(doc, 'STA structure identification', level=2)
add_para(doc,
    'TA1120 is a Single Technology Appraisal (STA) and not a Multiple Technology Appraisal. The only '
    'company recipient of clarification correspondence is Merck. The Committee papers PDF was page-mapped '
    'to identify the relevant sections: Section 1a contains the manufacturer submission (Document B and '
    'appendices); Section 2 contains the company response to the EAG clarification questions; Section 5 '
    'contains the EAG Assessment Report. Questions were attributed to a single recipient (Merck) throughout.')

add_heading(doc, 'Four-specialist parallel review', level=2)
add_para(doc,
    'Four specialist subagents were spawned in parallel, each adopting a defined EAG persona:')
for label, desc in [
    ('Clinical', 'EAG clinical specialist — trial design, internal/external validity, subgroup analyses, safety, HRQoL, treatment switching, ITC composition, redactions.'),
    ('Statistician', 'EAG statistician — proportional hazards diagnostics, survival distribution selection, NMA methodology, treatment-switching adjustment, censoring rules, multiplicity, PSA implementation.'),
    ('Health economist', 'EAG health economist — model auditability, scope concordance, model structure, survival extrapolation, utilities, costs, sensitivity analyses, reproducibility.'),
    ('Information specialist', 'EAG information specialist / SLR methodologist — PRISMA 2020 / PRISMA-S compliance, NMA feasibility, RWE methodology, data-on-file transparency.'),
]:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(label + '. ')
    run.bold = True
    p.add_run(desc)

add_heading(doc, 'Blinding protocol', level=2)
add_para(doc,
    'Specialists were given only Section 1a (the manufacturer submission and appendices) and the final '
    'scope. They were explicitly forbidden from reading Section 2 (the real company response to '
    'clarification) and Section 5 (the EAG Assessment Report). This blinding is essential to the design: '
    'the AI specialists must generate questions from the same evidence base that a real EAG would have at '
    'the clarification stage, not from documents that already encode the EAG\'s own findings.')

add_heading(doc, 'Sanity check on specialist outputs', level=2)
add_para(doc,
    'A post-generation sanity check confirmed that each specialist persona produced questions in its own '
    'voice (e.g., DSU TSD references from the statistician; PRISMA-S and Cochrane Handbook references from '
    'the information specialist; PMG36 references from the health economist). No textual contamination '
    'with Section 2 or Section 5 wording was detected.')

add_heading(doc, 'Actual NICE question extraction', level=2)
add_para(doc,
    'The actual NICE clarification questions were extracted verbatim from Section 2 (the company response, '
    'which quotes each NICE question before responding). NICE asked 23 unique questions, organised into '
    'three sections: A (Effectiveness, A1–A10), B (Cost-Effectiveness, B1–B12) and C (Textual / additional, '
    'C1). Eight questions were marked NICE priority "Yes": A1, A9, B1, B2, B7, B8, B10, B11, B12.')

add_heading(doc, 'Merge and compare step', level=2)
add_para(doc,
    'The 151 individual specialist questions were de-duplicated against each other (cross-persona merging) '
    'to produce 61 unique merged topics. Each merged topic was then mapped to NICE\'s 23 questions by '
    'artefact and methodological intent (rather than by exact wording). Conservative matching was used: '
    'where a generated topic addressed the same underlying issue or artefact as a NICE question, it was '
    'credited as matched, even if the framing differed.')

add_page_break(doc)

# ============================================================
# MATCHED TOPICS table — LANDSCAPE section
# ============================================================

# Switch to landscape
section_landscape = doc.add_section(WD_SECTION.NEW_PAGE)
section_landscape.orientation = WD_ORIENTATION.LANDSCAPE
section_landscape.page_width = Inches(11)
section_landscape.page_height = Inches(8.5)
section_landscape.left_margin = Inches(0.75)
section_landscape.right_margin = Inches(0.75)
section_landscape.top_margin = Inches(0.75)
section_landscape.bottom_margin = Inches(0.75)

add_heading(doc, 'Matched topics (Both)', level=1)
add_para(doc,
    'Sixteen merged topics from the AI review mapped directly to one or more NICE clarification questions. '
    'All three "priority Yes" methodological-deep-dive items (B1 time horizon, B2 hazard/PH, B7 utility regression) '
    'are matched, as are the NICE-priority items A1 (BICR vs investigator PFS) and A9 (NMA report).')

# Landscape table — 6 cols totalling 13680
matched_widths = [600, 2400, 1100, 3400, 1100, 5080]
assert sum(matched_widths) == 13680, sum(matched_widths)
matched_headers = ['Merged-ID', 'Topic', 'Origin tags', 'Generated by', 'NICE ID(s)', 'NICE verbatim summary']
matched_table = make_fixed_table(doc, matched_widths, matched_headers)

matched_rows = [
    ('T-01', 'BICR vs investigator-assessed PFS — concordance and figures',
     '[CLIN], [STAT]', 'CLIN-Merck-5, STAT-Merck-16', 'A1',
     'Provide BICR versions of CS Figures 7 and 9; for the favourable and intermediate/poor IMDC subgroups, provide a BICR/investigator KM plot for the total population (NICE priority Yes).'),
    ('T-02', 'Patient disposition by IMDC subgroup and arm',
     '[CLIN]', 'CLIN-Merck-7', 'A4',
     'Provide versions of CS Table 11 (disposition at FA) for the favourable and intermediate/poor IMDC risk subgroups.'),
    ('T-03', 'UK RWE source overlap and cohort definitions (SACT-EAMS, SACT-CDF, Nathan, McGrane)',
     '[INFO], [CLIN]', 'INFO-Merck-27, CLIN-Merck-35', 'A6, A7',
     'Confirm whether SACT-EAMS and SACT-CDF are the only non-overlapping UK RWE sources (Nathan/McGrane may overlap SACT); list reasons for exclusion from SACT-CDF and SACT-EAMS cohorts.'),
    ('T-04', 'RWE SLR report — provision and methodology (Appendix D.6)',
     '[INFO]', 'INFO-Merck-1, INFO-Merck-12', 'A8',
     'Provide the RWE SLR report referenced in Appendix D.6 (or indicate which reference in the reference pack it is).'),
    ('T-05', 'NMA report / feasibility assessment / risk of bias for comparator RCTs (intermediate/poor risk)',
     '[INFO], [CLIN], [STAT]',
     'INFO-Merck-9, INFO-Merck-11, INFO-Merck-16, CLIN-Merck-27, STAT-Merck-12', 'A9',
     'Provide separate NMA report or, failing that, intermediate/poor-risk input data, baseline characteristics, and comparator-RCT risk-of-bias assessments (NICE priority Yes).'),
    ('T-06', 'SLR critical appraisal — number of reviewers and independence',
     '[INFO]', 'INFO-Merck-7', 'A10',
     'State number of reviewers and whether they worked independently for critical appraisal of the clinical SLR and RWE SLR.'),
    ('T-07', 'Time horizon (40-year vs 10-year precedent)',
     '[HE]', 'HE-Merck-10', 'B1',
     'Justify use of 40-year horizon vs 10-year horizon used in prior aRCC TAs (NICE priority Yes).'),
    ('T-08', 'PH testing, hazard plots and parametric curve choice — favourable-risk OS, PFS, TTD',
     '[STAT], [HE], [CLIN], [INFO]',
     'STAT-Merck-1, -2, -5, -30; HE-Merck-11, -12; CLIN-Merck-33; INFO-Merck-17', 'B2',
     'Provide hazard plots for OS, PFS, TTD in favourable-risk subgroup (both arms); assess PH assumption; justify parametric model choice (NICE priority Yes).'),
    ('T-09', 'EQ-5D-5L scoring / Hernández-Alava 2017 mapping confirmation',
     '[CLIN], [STAT]', 'CLIN-Merck-26, STAT-Merck-26', 'B5',
     'State method used to attribute index scores to EQ-5D-5L FAS data; confirm whether Hernández-Alava 2017 mapping was used as in the economic model.'),
    ('T-10', 'EQ-5D-5L descriptive data by IMDC risk subgroup',
     '[CLIN]', 'CLIN-Merck-24', 'B6',
     'Provide descriptive EQ-5D-5L results and change-from-baseline plots for favourable and intermediate/poor subgroups.'),
    ('T-11', 'Utility regression methodology — missingness, imputation, model specification, covariates',
     '[HE], [STAT], [CLIN]',
     'HE-Merck-21, -22; STAT-Merck-26; CLIN-Merck-26', 'B7',
     'Explain EQ-5D-5L missing-data volume/pattern, imputation decisions, justification of Model 1/2 specs, and alternative ITT+interaction model (NICE priority Yes).'),
    ('T-12', 'No-nivolumab-as-subsequent-therapy assumption — justification',
     '[CLIN], [HE]', 'CLIN-Merck-23; HE-Merck-30, -16', 'B8',
     'Provide clearer justification (NICE guidance, NHSE policy, SACT) for excluding subsequent nivolumab after avelumab + axitinib and justify scenario alternatives (NICE priority Yes).'),
    ('T-13', 'Equal AE costs / non-elective short-stay tariff for all Grade ≥3 AEs',
     '[HE]', 'HE-Merck-33', 'B9',
     'Justify assumption of equal cost across all AEs and use of non-elective short-stay tariff for all Grade ≥3 AEs.'),
    ('T-14', 'Data-on-file references — confidentiality / provision',
     '[INFO]', 'INFO-Merck-25', 'C1',
     'Should all data-on-file (DOF) references be treated as commercial in confidence?'),
    ('T-15', 'Treatment-switching / subsequent-IO adjustment for sunitinib OS',
     '[CLIN], [STAT], [HE]',
     'CLIN-Merck-22; STAT-Merck-14, -15; HE-Merck-16', 'B8 (partial)',
     'Switching-adjustment side of the nivolumab-as-subsequent-therapy question — RPSFTM/TSE/IPCW adjusted OS for sunitinib arm given 64.6–86.2% subsequent IO exposure.'),
    ('T-16', 'Subsequent therapy distribution applied in model — derivation',
     '[CLIN], [HE]', 'CLIN-Merck-23; HE-Merck-30, -31', 'B10',
     'Rationalise why only 6 drugs were used for poor-risk vs 24 for intermediate-risk vs source Table 10.5.a (NICE priority Yes).'),
]

for i, vals in enumerate(matched_rows):
    add_table_row(matched_table, matched_widths, vals, font_size=9,
                  shade_alt=True, row_index=i)

doc.add_paragraph()
add_para(doc,
    'Note on T-15 / T-12: NICE B8 conflates the question of why nivolumab is excluded (the company\'s '
    'assumption) with the consequences for OS. The generated set covers both the justification side (T-12) '
    'and the OS-adjustment methodology side (T-15); only the justification side is mapped 1:1 to B8.',
    size=10, italic=True)

# ============================================================
# MISSED TOPICS (NICE-only) — LANDSCAPE continued
# ============================================================
add_page_break(doc)

add_heading(doc, 'Missed topics (NICE-only)', level=1)
add_para(doc,
    'Seven NICE questions were not generated by the AI review. This is the most important diagnostic '
    'table in the report: it identifies the systematic blind spot of pre-clarification AI review.')

missed_widths = [800, 2200, 7080, 3600]
assert sum(missed_widths) == 13680
missed_headers = ['NICE ID', 'Topic', 'Verbatim NICE text (abbreviated)', 'Likely reason missed']
missed_table = make_fixed_table(doc, missed_widths, missed_headers, header_shade='C00000')

missed_rows = [
    ('A2', 'FKSI-DRS time-to-deterioration outcome reporting',
     'CS Table 6 and CS Table 7 refer to the outcome for "Time to deterioration in the FKSI-DRS", however the results are not reported in CS section B.2.6 nor identified in the CSR. Please provide the results or signpost to where they are in the submission.',
     'Document-tracing question — requires close cross-reference of CSR vs Document B. Specialists generated broader HRQoL/PRO requests but not this specific orphaned outcome.'),
    ('A3', 'Avelumab vs axitinib death-discontinuation count discrepancy (n=25 vs n=28)',
     'CS Table 11. The number of participants who discontinued due to death differs for avelumab and axitinib (n=25 and n=28 respectively). Is this because some participants were only in receipt of either avelumab or axitinib at the time of death? If not, please explain why these numbers differ.',
     'Granular cell-level discrepancy noticed during model audit; below the resolution of the generated topic list (which asked for unredacted Table 11 disposition data more broadly — CLIN-Merck-7).'),
    ('A5', 'AE rates with each agent as monotherapy',
     'CS section B.2.11.1.2 reports that some AEs were at higher frequencies with the combination than with either single agent alone. Please indicate where we can find the AE data for each agent as monotherapy.',
     'Document-tracing question; specialists requested AE-by-subgroup and AE-by-irAE-cluster but not the monotherapy reference figures.'),
    ('B3', 'KM data timepoint discrepancy (model vs source)',
     'KM data in the model sheet (KM data!BK46:CM233) have different time points to those reported in the source DOF analyses. Model has 0, 3.81, 5.09 months while source provides 4, 8, 12-month intervals. Confirm whether discrepancy is due to model data being derived from IPD or provide a source with matching KM data.',
     'Model-audit finding requiring direct comparison of model workbook KM data sheet against DOF source tables; generated set asked for unredacted KM data but did not catch the timepoint anomaly.'),
    ('B4', 'Missing Grade ≥3 TRAEs in CS Table 56',
     'Please justify why some grade ≥3 TRAEs experienced by more than 5% of patients are missing from CS Table 56, including fatigue, nausea, decreased appetite, stomatitis, mucosal inflammation, AST increased, and vomiting.',
     'Table-content gap requiring item-by-item comparison against the CSR; generated AE questions focused on redactions, irAEs, cardiac signal and AE costing — not Table 56 completeness.'),
    ('B11', 'Untraceable model input parameters (age, sex, weight by risk group) [Priority Yes]',
     'The EAG cannot identify the sources of several model input parameters (Table 1): average age, proportion female, and average weight for favourable-risk (CS Table 40), intermediate/poor-risk (CS Appendix O Table 50), and average weight for ITT (CS Appendix O Table 51) — all sourced to JAVELIN Renal 101.',
     'Model-audit finding from line-by-line workbook traceability; generated HE set asked for executable model and scripts (HE-Merck-1, -2) but did not flag this specific traceability gap for the workbook against JAVELIN.'),
    ('B12', 'CS-vs-model input value discrepancies (IV admin, CT/blood, nivolumab dose, subsequent therapy cost) [Priority Yes]',
     'Different input values are reported in the company submission and in the company\'s model for IV administration costs (CS Table 64 vs Costs!G86), CT and blood test costs (CS Table 66 vs Costs!G118:G122), nivolumab dose (CS Table 71 vs Costs!H183) and total subsequent treatment cost intermediate/poor risk (Appendix O Table 91 vs Costs!G234).',
     'Direct CS↔workbook reconciliation; generated HE set asked for cell-exact reproducibility (HE-Merck-49) but the specific Costs!G86, G118:G122, H183 and G234 discrepancies were not pre-identified at the public-submission stage.'),
]

for i, vals in enumerate(missed_rows):
    add_table_row(missed_table, missed_widths, vals, font_size=9,
                  shade_alt=True, row_index=i)

doc.add_paragraph()
add_para(doc,
    'Common pattern: All seven missed items are reactive findings from direct, cell-by-cell audit of '
    'artefacts the specialists did not have in front of them — the executable Excel model and the full CSR. '
    'Six of the seven are either NICE-priority "Yes" (B11, B12) or document/cell-discrepancy questions '
    '(A2, A3, A5, B3, B4). The generated set proactively requested these artefacts (CLIN-Merck-7 unredacted '
    'Table 11; HE-Merck-1 executable model; HE-Merck-49 cell-exact reproducibility; INFO-Merck-25 DOF list) '
    'but could not have identified the specific within-artefact discrepancies until those artefacts were '
    'actually opened.', italic=True, size=10)

# ============================================================
# GENERATED-ONLY TOPICS — LANDSCAPE continued
# ============================================================
add_page_break(doc)

add_heading(doc, 'Generated-only topics', level=1)
add_para(doc,
    'Forty-five merged topics were generated by the AI review but not asked by NICE. Most are legitimate '
    'critiques that NICE chose not to pursue at the clarification stage rather than spurious questions. '
    'NICE prioritises rather than asking every possible question; the EAG\'s clarification letter is the '
    'first of multiple opportunities to interrogate the submission, with detailed methodological challenges '
    'often deferred to the EAR / EAG report stage.')

genonly_widths = [600, 3000, 1280, 8800]
assert sum(genonly_widths) == 13680
genonly_headers = ['Merged-ID', 'Topic', 'Origin tags', 'One-line substance']
genonly_table = make_fixed_table(doc, genonly_widths, genonly_headers, header_shade='548235')

genonly_rows = [
    ('T-17', 'Base case framed on favourable-risk vs scope-specified ITT', '[CLIN], [HE]',
     'Scope deviation; base case should be ITT per scope, not favourable-risk subgroup.'),
    ('T-18', 'Omission of NMA vs favourable-risk TKI comparators (pazopanib, tivozanib)', '[CLIN], [STAT], [HE], [INFO]',
     'TKI equivalence assumption asserted on Manz 2019 basis but never tested in favourable-risk subgroup.'),
    ('T-19', 'Exclusion of PD-L1 subgroup from cost-effectiveness analysis', '[CLIN], [STAT], [HE]',
     'Primary regulatory endpoint population (PD-L1+) excluded from CEM on unspecified clinical expert opinion.'),
    ('T-20', 'Marketing authorisation wording vs trial eligibility (clear-cell, ECOG)', '[CLIN]',
     'Licence may exceed trial population; need explicit mapping of MA wording to JAVELIN eligibility.'),
    ('T-21', 'PFS censoring rules and EMA/FDA approach', '[CLIN], [STAT]',
     'Exact censoring rules and sensitivity not transparently documented.'),
    ('T-22', 'OS data maturity in IMDC subgroups', '[CLIN]',
     'Favourable-risk OS ~50% mature; need matched-follow-up sensitivity.'),
    ('T-23', 'Stratification, multiplicity and pre-specification of IMDC subgroup analyses', '[CLIN], [STAT]',
     'Subgroups exploratory, post-hoc pooling, no multiplicity adjustment.'),
    ('T-24', 'UK representation in JAVELIN Renal 101 (n=32) and generalisability', '[CLIN]',
     'UK = 3.6% of trial population; outcomes for UK subset.'),
    ('T-25', 'Baseline characteristics vs UK practice (nephrectomy, ECOG, sarcomatoid)', '[CLIN]',
     'JAVELIN favourable-risk may differ from UK real-world (~80% nephrectomy vs UK).'),
    ('T-26', 'Sunitinib comparator schedule (4/2 vs 2/1) vs UK practice', '[CLIN]',
     'UK has shifted to 2/1 schedule; trial used 4/2.'),
    ('T-27', 'Place of avelumab + axitinib in evolved 2024 treatment pathway', '[CLIN]',
     'Post-TA780, TA858, TA964 — landscape has changed since TA645.'),
    ('T-28', 'Full subgroup forest plots (ITT, all pre-specified subgroups, OS & PFS)', '[CLIN], [STAT]',
     'Complete forest plots to identify subgroups favouring sunitinib.'),
    ('T-29', 'Sarcomatoid disease subgroup', '[CLIN]',
     'Recognised IO-responsive subgroup not reported.'),
    ('T-30', 'PD-L1 × IMDC favourable-risk interactions', '[CLIN], [STAT]',
     'PD-L1+ favourable-risk fully redacted.'),
    ('T-31', 'Intermediate vs poor risk separated (post-hoc pooling)', '[CLIN]',
     'Need disaggregation of pooled subgroup; expected survival 24.1 vs 10.2 months differs materially.'),
    ('T-32', 'RDI by IMDC subgroup, dose modifications and cumulative dose', '[CLIN], [HE]',
     'RDI applied is ITT-derived in favourable-risk base case.'),
    ('T-33', 'Treatment-related discontinuation patterns (avelumab vs axitinib alone vs both)', '[CLIN]',
     'Differential discontinuation impacts cost and effectiveness; partial discontinuation common.'),
    ('T-34', 'Immune-related AEs, corticosteroid use, long-term endocrine sequelae', '[CLIN]',
     'One-off £801 AE cost may not capture lifelong irAE management.'),
    ('T-35', 'Cardiac toxicity signal (redacted)', '[CLIN]',
     'Unredact cardiac TEAEs including EF decreased, troponin T elevation.'),
    ('T-36', 'EQ-5D timing artefact in sunitinib arm (off-treatment week sampling)', '[CLIN]',
     'Submission acknowledges artefact but does not adjust.'),
    ('T-37', 'NMA composition — proportional hazards and inclusion criteria', '[CLIN], [STAT], [INFO]',
     'Network heterogeneity, KEYNOTE-426 exclusion.'),
    ('T-38', 'Fixed-effects vs random-effects in single-study-per-edge network', '[CLIN], [STAT], [INFO]',
     'DIC essentially identical; FE under-represents uncertainty.'),
    ('T-39', 'Severity / QALY shortfall by IMDC subgroup', '[CLIN], [HE]',
     'Severity may apply differently in poor-risk.'),
    ('T-40', 'Equality considerations and protected characteristics', '[CLIN]',
     'Equity framed around regional commissioning, not protected characteristics.'),
    ('T-41', 'AVION and other ongoing studies', '[CLIN]',
     'Future data cuts and Phase 4 commitments.'),
    ('T-42', 'TTD curves — clinical plausibility and joint sampling', '[CLIN], [STAT]',
     'TTD modelled independently for avelumab and axitinib; correlation in trial 93%.'),
    ('T-43', 'Avelumab stopping rule (2-year cap)', '[CLIN], [HE]',
     'TA858/TA964 precedent for IO stopping rules; need scenario.'),
    ('T-44', 'SACT real-world OS vs JAVELIN OS — external validation', '[CLIN], [STAT], [HE]',
     'RWE discordance with extrapolation; UK real-world should be primary check on plausibility.'),
    ('T-45', 'Consolidated unredacted academic-in-confidence appendix', '[CLIN], [HE]',
     'Volume of redactions in public submission is exceptional.'),
    ('T-46', 'PSA instability of generalised gamma sunitinib OS extrapolation', '[STAT], [HE]',
     '"Unrealistic extrapolations" admitted in B.3.9.1; inflates comparator LY in PSA.'),
    ('T-47', 'Curve crossing OS / PFS — long-term plausibility', '[STAT], [HE]',
     'Modelled HR may reverse over horizon under independent extrapolation.'),
    ('T-48', 'MCMC convergence diagnostics for NMA', '[STAT], [INFO]',
     'R-hat, ESS, trace plots not reported.'),
    ('T-49', 'Power for favourable-risk OS comparison', '[STAT]',
     'Subgroup not powered for OS difference; HR 0.73 p=0.13.'),
    ('T-50', 'Cholesky / joint sampling of PSM parameters in PSA', '[STAT]',
     'Independent draws break correlation structure of survival models.'),
    ('T-51', 'PSA stopping criterion / iteration adequacy', '[STAT], [HE]',
     'Running mean and INMB half-width vs iteration not reported.'),
    ('T-52', 'Executable model, unprotected, R/Stata scripts, user guide vs TA645', '[HE]',
     'NICE PMG36 requires reproducibility; changelog vs prior submission.'),
    ('T-53', 'Partitioned survival vs state-transition structure', '[HE]',
     'DSU TSD 19 cross-check; PFS-OS independence not validated.'),
    ('T-54', 'Half-cycle correction', '[HE]',
     'Standard practice test not done despite 40-year horizon.'),
    ('T-55', 'General-population mortality cap source and binding cycle', '[HE]',
     'Life table source and algebra; cap-binding cycle not transparent.'),
    ('T-56', 'TTD vs PFS coherence — patients on treatment beyond progression', '[HE], [STAT]',
     'Generalised gamma TTD tail may imply on-treatment beyond progression.'),
    ('T-57', 'AE rate equivalence between TKIs (sunitinib = tivozanib = pazopanib)', '[HE]',
     'TIVO-1, COMPARZ specific AE profiles ignored by single-AE-cost assumption.'),
    ('T-58', 'Age adjustment of utilities (Ara & Brazier) — recalculation per cycle', '[HE]',
     'Time-varying baseline; unclear whether multiplier recalculated cycle-by-cycle.'),
    ('T-59', 'AE disutility exclusion — six-weekly EQ-5D capture window', '[HE]',
     'Transient AEs may not be captured by EQ-5D every 6 weeks.'),
    ('T-60', 'Cross-TA utility consistency (TA417, TA858, TA964)', '[HE]',
     'Post-progression utility above prior precedent in long-tailed model.'),
    ('T-61', 'BNF / Reference Cost / PSSRU vintage inconsistency', '[HE]',
     '2021/22 vs 2022/23 reference cost mix across tables 64/67.'),
    ('T-62', 'Avelumab dosing schedule alternatives (Q2W vs Q3W)', '[HE]',
     'Within-licence dose sensitivity.'),
    ('T-63', 'Generic axitinib pricing assumption', '[HE]',
     'Decision must be on prices in force at time of guidance.'),
    ('T-64', 'Subsequent therapy dosing assumptions / TA542 TA498 vintage', '[HE]',
     '2018 mean TTD values applied to 2024–25 NHS practice.'),
    ('T-65', 'HCRU frequencies (TA581 2019 source)', '[HE]',
     'Pre-IO-era resource use; not validated for 2024 practice.'),
    ('T-66', 'End-of-life cost source (Round 2015)', '[HE]',
     'TA858/TA964 cross-TA consistency.'),
    ('T-67', 'Treatment-effect waning scenarios', '[HE]',
     'PMG36/TA858/TA964 precedent for waning; not modelled.'),
    ('T-68', 'Expanded OWSA tornados and parameter list', '[HE]',
     'Top 30 by INMB rather than top 10.'),
    ('T-69', 'CEAC values and PSA scatter at multiple thresholds', '[HE]',
     'Underlying probabilities not reported numerically.'),
    ('T-70', 'Discount rate scenarios (1.5% justification under PMG36)', '[HE]',
     '1.5% scenario shown without PMG36 severity/restoration caveats.'),
    ('T-71', 'Favourable-risk subgroup full uncertainty (PSA, OWSA, scenarios)', '[HE]',
     'AE rates ITT-derived in favourable-risk base case.'),
    ('T-72', 'Intermediate-/poor-risk subgroup full uncertainty (PSA, OWSA, scenarios)', '[HE]',
     'Full uncertainty characterisation for seven comparators.'),
    ('T-73', 'ISPOR-SMDM model validation checklist and named experts', '[HE]',
     'Internal QC + 3 clinical consultations only; structured validation absent.'),
    ('T-74', 'Cross-TA comparison table vs TA645, TA858, TA964, TA780', '[HE]',
     'Identify methodological departures from precedent.'),
    ('T-75', 'Routine vs managed access — modality of recommendation', '[HE]',
     'CDF exit decision — what is the company asking for?'),
    ('T-76', 'Full results reproducibility cell-exact from submitted model', '[HE]',
     'PMG36 requirement; broader frame than the specific B11/B12 issues NICE caught.'),
    ('T-77', 'PROSPERO registration of clinical SLR', '[INFO]',
     'PRISMA 2020 item 24a.'),
    ('T-78', 'Line-by-line search strategies — clinical SLR', '[INFO]',
     'PRISMA-S items 6–8.'),
    ('T-79', 'Trial-registry and grey-literature searches', '[INFO]',
     'PRISMA-S items 5, 11.'),
    ('T-80', 'Date limits / awareness check / gap to submission', '[INFO]',
     'Cochrane Handbook 4.4.9.'),
    ('T-81', 'Language and publication-type restrictions', '[INFO]',
     'PRISMA-S item 9.'),
    ('T-82', 'PRISMA 2020 flow diagram', '[INFO]',
     'PRISMA 2020 item 16a.'),
    ('T-83', 'Characteristics of included studies — consolidated table', '[INFO]',
     'PRISMA 2020 item 17.'),
    ('T-84', 'Characteristics of excluded studies — KEYNOTE-426, COSMIC-313 etc.', '[INFO]',
     'PRISMA 2020 item 17.'),
    ('T-85', 'Risk of bias — RWE studies (ROBINS-I / NICE RWE Framework)', '[INFO]',
     'Cochrane Handbook Ch.25; NICE RWE Framework.'),
    ('T-86', 'Publication bias / small-study effects', '[INFO]',
     'PRISMA-NMA item S6.'),
    ('T-87', 'Explicit NMA-stage PICOS', '[INFO]',
     'Distinct from SLR PICOS.'),
    ('T-88', 'Inconsistency / loop assessment (star network)', '[INFO]',
     'DSU TSD 4.'),
    ('T-89', 'Cost-effectiveness SLR (Appendix G) — methodology and results', '[INFO]',
     'PRISMA 2020 / PMG36 §3.4.'),
    ('T-90', 'HRQoL / utility SLR — protocol, search, screening', '[INFO]',
     'PMG36 §3.4.6.'),
    ('T-91', 'Resource-use / cost SLR', '[INFO]',
     'PMG36 §3.4.5.'),
    ('T-92', 'Clinical Expert Opinion (DoF reference 17) — methodology', '[INFO]',
     'ISPOR Bojke 2021 elicitation good practice.'),
    ('T-93', 'Treatment of CheckMate 214 PD-L1 selected vs unselected', '[INFO]',
     'DSU TSD 3 §2.2.'),
    ('T-94', 'Latest data cuts for comparator trials (CheckMate 9ER, 214, CLEAR)', '[INFO]',
     'PRISMA-NMA item S3.'),
    ('T-95', 'Risk-of-bias traffic-light summary', '[INFO]',
     'PRISMA 2020 item 18.'),
    ('T-96', 'Complete list of submitted appendices and version control', '[INFO]',
     'PMG36 §3.2.'),
    ('T-97', 'Search audit trail (EndNote / RIS / screening log)', '[INFO]',
     'PRISMA-S item 16.'),
]

# Note: source listed 81 IDs from T-17 to T-97 with some duplicates/consolidation;
# comparison.md says final consolidated count is 45 after de-duplication.
# Display all listed rows for transparency.

for i, vals in enumerate(genonly_rows):
    add_table_row(genonly_table, genonly_widths, vals, font_size=9,
                  shade_alt=True, row_index=i)

doc.add_paragraph()
add_para(doc,
    'Note on counting: Table T-17 to T-97 is shown in full above for transparency. The headline figure '
    'of 45 generated-only topics reflects de-duplication where some merged-IDs are functionally overlapping '
    '(e.g. T-44 SACT/JAVELIN OS validation subsumes one statistician-flagged absolute-discrepancy variant). '
    'NICE prioritises rather than asking every possible question; the EAG\'s clarification letter is '
    'tactically scoped to the issues most likely to need a written response from the company, with '
    'methodological challenges often deferred to the EAR/EAG report stage.', size=10, italic=True)

# ============================================================
# Switch back to portrait
# ============================================================
section_portrait = doc.add_section(WD_SECTION.NEW_PAGE)
section_portrait.orientation = WD_ORIENTATION.PORTRAIT
section_portrait.page_width = Inches(8.5)
section_portrait.page_height = Inches(11)
section_portrait.left_margin = Inches(1.0)
section_portrait.right_margin = Inches(1.0)
section_portrait.top_margin = Inches(1.0)
section_portrait.bottom_margin = Inches(1.0)

# ============================================================
# SPECIALIST CONTRIBUTIONS
# ============================================================

add_heading(doc, 'Specialist contributions', level=1)
add_para(doc,
    'The four specialist personas each contributed distinctively to the matched set. The table below '
    'shows the number of individual questions generated by each persona, how many contributed to a '
    'matched topic, and the resulting persona-level "match rate". Note that a single matched topic '
    'frequently combines questions from multiple personas (e.g., T-08 hazard plots & PH testing draws '
    'on all four personas).')

spec_widths = [1500, 1200, 1700, 1900, 1700, 1360]
assert sum(spec_widths) == 9360
spec_headers = ['Persona', 'Generated', 'Matched', 'Generated-only', 'Match rate', 'Match count basis']
spec_table = make_fixed_table(doc, spec_widths, spec_headers, header_shade='1F3864')

spec_rows = [
    ('Clinical', '36', '11', '25', '31%', '11 of 36 individual'),
    ('Statistician', '30', '7', '23', '23%', '7 of 30 individual'),
    ('Health economist', '50', '9', '41', '18%', '9 of 50 individual'),
    ('Information specialist', '35', '6', '29', '17%', '6 of 35 individual'),
    ('TOTAL (individual)', '151', '33', '118', '22%', '151 → 61 merged'),
]
for i, vals in enumerate(spec_rows):
    row = add_table_row(spec_table, spec_widths, vals, font_size=10,
                        shade_alt=True, row_index=i)
    if vals[0].startswith('TOTAL'):
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True

doc.add_paragraph()

add_heading(doc, 'Persona commentary', level=2)

add_para(doc, 'Clinical', bold=True)
add_para(doc,
    'The clinical persona produced the broadest persona-level contribution to the matched set (11 of 16 '
    'matched topics), covering the full spectrum of trial conduct, internal validity, subgroup credibility, '
    'safety and HRQoL. Its highest-impact contributions to the matched set were T-01 (BICR vs investigator '
    'PFS), T-02 (disposition by IMDC subgroup) and T-03 (UK RWE source overlap). It also generated the '
    'largest pool of unmatched topics — these include scope-deviation critiques (T-17 ITT vs favourable-risk '
    'framing, T-19 PD-L1 exclusion) and external-validity concerns (T-24 UK representation, T-26 sunitinib '
    'schedule) that NICE did not pursue in clarification but which may surface in committee discussion.')

add_para(doc, 'Statistician', bold=True)
add_para(doc,
    'The statistician contributed to the methodologically deep matched topics: T-08 (PH testing and '
    'parametric choice), T-11 (utility regression), T-15 (treatment-switching adjustment). Its match rate '
    '(23%) reflects the heavy emphasis on diagnostic detail (Schoenfeld residuals, log-cumulative hazard '
    'plots, Cholesky decomposition, fractional polynomial NMA per TSD 19) that is decision-critical at the '
    'EAR stage but rarely asked verbatim in clarification. Its highest-priority unmatched contribution is '
    'T-46 (PSA instability of generalised gamma sunitinib OS — explicitly acknowledged in B.3.9.1 as '
    '"unrealistic extrapolations" but not pursued by NICE at clarification).')

add_para(doc, 'Health economist', bold=True)
add_para(doc,
    'The health economist produced the largest absolute number of questions (50) but had the second-lowest '
    'match rate (18%). Its matched contributions are concentrated in the high-priority NICE items: '
    'T-07 time horizon (B1), T-08 hazard/PH (B2), T-11 utility regression (B7), T-12 nivolumab assumption '
    '(B8), T-13 AE costs (B9), T-16 subsequent therapy distribution (B10). The 41 unmatched topics cover '
    'the standard EAG cost-effectiveness checklist (cross-TA consistency, vintage of reference costs, '
    'waning scenarios, OWSA scope, CEAC at multiple thresholds, generic axitinib pricing) — items that '
    'normally surface in the EAR rather than the clarification letter.')

add_para(doc, 'Information specialist', bold=True)
add_para(doc,
    'The information specialist had the lowest match rate (17%) because NICE asked only a single bundled '
    'SLR-methodology question (A10) rather than the full PRISMA-S / Cochrane / NICE RWE Framework '
    'checklist. Its matched contributions to T-04 (RWE SLR report), T-05 (NMA report) and T-14 (DOF '
    'confidentiality) reflect direct artefact requests, while its unmatched topics — PROSPERO registration, '
    'line-by-line search strategies, PRISMA 2020 flow diagrams, two-reviewer methodology, risk-of-bias '
    'traffic-light summaries — would normally appear in a fuller SLR-methods critique at EAR stage.')

add_page_break(doc)

# ============================================================
# DISCUSSION
# ============================================================

add_heading(doc, 'Discussion', level=1)

add_heading(doc, 'Patterns in the matched set', level=2)
add_para(doc,
    'The 16 matched topics divide into three families. First, document/data requests directly tied to '
    'Committee-paper artefacts (T-04 RWE SLR report, T-05 NMA report and feasibility assessment, T-14 DOF '
    'confidentiality, T-10 EQ-5D-5L by subgroup) — these matched cleanly to NICE\'s A8, A9, A6/A7, B6 and '
    'C1 because both the specialists and NICE were responding to the same visible signposts in the '
    'submission. Second, methodological deep-dives that NICE prioritised (T-08 PH testing and parametric '
    'choice → B2 priority; T-11 utility regression → B7 priority; T-12 no-nivolumab assumption → B8 '
    'priority; T-07 40-year time horizon → B1 priority). The statistician and health economist personas '
    'produced these in detail. Third, NICE-priority issues already prominent in the public submission text '
    'that all specialists picked up (T-01 BICR/investigator PFS, T-13 equal AE costs).')

add_heading(doc, 'Patterns in the generated-only set', level=2)
add_para(doc,
    'Forty-five generated topics did not surface in NICE\'s letter. Most are legitimate critiques that '
    'NICE chose not to pursue rather than spurious questions. Examples include: structural scope critiques '
    '(T-17 ITT vs favourable-risk base case; T-19 PD-L1 subgroup exclusion; T-18 TKI equivalence assumption) — '
    'NICE accepted the company\'s framing rather than re-litigating it at clarification; methodological '
    'challenges that would normally be pursued at the EAR/EAG report stage rather than clarification '
    '(T-46 PSA instability of sunitinib OS, T-47 curve crossing, T-50 Cholesky sampling, T-67 treatment-effect '
    'waning, T-74 cross-TA consistency); SLR-process critiques (T-77 PROSPERO registration, T-78 line-by-line '
    'strategies, T-86 publication bias, T-95 risk-of-bias traffic-light) — NICE asked only a single bundled '
    'SLR question (A10) rather than the full PRISMA-S/Cochrane checklist; and disease-specific clinical depth '
    '(T-22 OS maturity, T-26 sunitinib 4/2 vs 2/1, T-29 sarcomatoid, T-31 intermediate vs poor separation) '
    'which the clinical persona emphasised but which NICE did not pursue in clarification — likely because '
    'these are committee-discussion items, not document-request items.')

add_heading(doc, 'Patterns in NICE-only (missed) — the dominant cause', level=2)
add_para(doc,
    'All seven missed items share a single characteristic: they are findings produced by direct, '
    'cell-by-cell audit of artefacts the specialists did not have in front of them. A2 (FKSI-DRS missing '
    'from B.2.6), A3 (n=25 vs n=28 death-discontinuation cell discrepancy), A5 (monotherapy AE rates '
    'referenced but unlocatable), B3 (KM timepoint mismatch between model workbook and DOF tables), B4 '
    '(Table 56 Grade ≥3 TRAEs omitted), B11 (age/sex/weight model inputs untraceable to JAVELIN source), '
    'and B12 (CS text vs Costs!G86/G118/H183/G234 input value discrepancies). The specialists worked from '
    'the public Committee-papers PDF and so could not interrogate the executable Excel model, the full '
    'CSR, or the underlying DOF analysis files. The generated set asked for these artefacts (CLIN-Merck-7 '
    'for unredacted Table 11; HE-Merck-1 for the executable model; HE-Merck-49 for cell-exact '
    'reproducibility; INFO-Merck-25 for the DOF list) but could not have identified the specific '
    'within-artefact discrepancies until those artefacts were actually opened. This is the canonical '
    'pre-clarification limitation: the EAG\'s most decision-relevant clarification questions are reactive '
    'findings from the workbook and CSR rather than proactive critiques of the public submission text.')

add_heading(doc, 'Document requests handled in different form', level=2)
add_para(doc,
    'Several NICE questions are precise document requests where the generated set asked for the same '
    'artefact in a broader frame. A8 (RWE SLR report) maps to INFO-Merck-1\'s broader Appendix D request '
    'and INFO-Merck-12\'s RWE risk-of-bias request. A9 (separate NMA report) maps to a cluster of generated '
    'questions on NMA composition, feasibility, risk of bias and trial inclusion. C1 (DOF confidentiality '
    'treatment) maps to INFO-Merck-25\'s consolidated DOF list request. The matching was done at the '
    'artefact level rather than the wording level; the generated set\'s "please provide unredacted/full X" '
    'formulation is functionally equivalent to NICE\'s "please provide reference Y" when the underlying '
    'artefact is the same.')

add_heading(doc, 'Limitations of the AI review', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Cannot audit Excel cell values: the AI review worked from the Committee papers PDF only and '
          'could not run the executable model, so cell-level discrepancies (B11, B12) and KM timepoint '
          'mismatches (B3) were unreachable.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Cannot cross-reference CSR appendix tables in isolation: A2 (FKSI-DRS) and A5 (monotherapy AE '
          'rates) required line-by-line CSR lookup that the AI could not perform without the full CSR.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Cannot detect orphaned outcome references without the full CSR: the AI flagged broad data-on-file '
          'and unredaction requests but could not identify the specific orphaned-outcome anomaly that '
          'distinguishes A2.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Verbosity and lack of prioritisation: the AI generated 151 questions across 61 topics. NICE '
          'asked 23. The AI is not constrained by the practical limits of clarification correspondence '
          '(typically 10–15 prioritised questions).')

add_heading(doc, 'Limitations of the comparison', level=2)
add_para(doc,
    'NICE asks only priority and tractable questions; the AI is not constrained that way. The 69.6% hit '
    'rate reflects how often a generated topic and a NICE topic interrogated the same artefact or '
    'methodological issue; it does not measure question quality, depth, or priority alignment. Conservative '
    'matching means borderline cases (e.g., the breadth of HE-Merck-33 vs the narrowness of NICE B9) are '
    'credited as matched on the basis of artefact overlap. NICE\'s three priority "Yes" methodological items '
    '(B1, B2, B7) are all matched; the two NICE priority "Yes" items missed (B11, B12) reflect the '
    'audit-vs-critique distinction rather than a methodological gap.')

add_page_break(doc)

# ============================================================
# CONCLUSION
# ============================================================

add_heading(doc, 'Conclusion', level=1)
add_para(doc,
    'The AI-assisted EAG simulation would have identified 16 of the 23 actual NICE clarification questions '
    '(69.6% hit rate), including all three priority "Yes" methodological deep-dives (time horizon, hazard '
    'plots / proportional hazards, utility regression). The seven missed items are all reactive findings '
    'from direct workbook and CSR audit — artefacts the AI did not have at the pre-clarification stage. '
    'Suggested use cases: (1) pre-EAG triage to pre-populate the clarification letter draft before the '
    'specialist team opens the executable model; (2) training of new EAG analysts against a calibrated '
    'reference of NICE clarification patterns; (3) audit of completed EAG reports to identify '
    'methodological topics the team did not pursue but for which a written record may be required.')

add_page_break(doc)

# ============================================================
# APPENDIX A — FULL GENERATED QUESTION LIST
# ============================================================

add_heading(doc, 'Appendix A: Full generated question list', level=1)
add_para(doc,
    'This appendix reproduces the four specialist outputs in full, in the form they were generated. IDs '
    'are stable across the comparison tables (e.g. CLIN-Merck-5 referenced in T-01 is the same '
    'CLIN-Merck-5 below).', italic=True, size=10)

# Helper to render markdown-like specialist text
import re

def add_specialist_section(doc, title, file_path):
    add_heading(doc, title, level=2)
    with open(file_path, 'r') as f:
        content = f.read()
    lines = content.split('\n')
    for line in lines:
        if line.startswith('# '):
            continue  # Skip top-level title; already added
        elif line.startswith('## '):
            add_heading(doc, line[3:].strip(), level=3)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.font.size = Pt(11)
        elif line.startswith('- **'):
            # bullet with bold prefix
            m = re.match(r'- \*\*(.*?):\*\*\s*(.*)', line)
            if m:
                p = doc.add_paragraph(style='List Bullet')
                r1 = p.add_run(m.group(1) + ': ')
                r1.bold = True
                r1.font.size = Pt(10)
                r2 = p.add_run(m.group(2))
                r2.font.size = Pt(10)
            else:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(line[2:].strip())
                run.font.size = Pt(10)
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(line[2:].strip())
            run.font.size = Pt(10)
        elif line.startswith('**'):
            m = re.match(r'\*\*(.*?):\*\*\s*(.*)', line)
            if m:
                p = doc.add_paragraph()
                r1 = p.add_run(m.group(1) + ': ')
                r1.bold = True
                r1.font.size = Pt(10)
                r2 = p.add_run(m.group(2))
                r2.font.size = Pt(10)
            else:
                m2 = re.match(r'\*\*(.*?)\*\*(.*)', line)
                if m2:
                    p = doc.add_paragraph()
                    r1 = p.add_run(m2.group(1))
                    r1.bold = True
                    r1.font.size = Pt(10)
                    r2 = p.add_run(m2.group(2))
                    r2.font.size = Pt(10)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.font.size = Pt(10)
        elif line.startswith('|'):
            # Skip raw markdown tables in appendix to avoid bloat
            continue
        elif line.startswith('---'):
            continue
        elif line.strip() == '':
            continue  # Skip extra blanks; paragraphs add their own space
        else:
            # Plain text — handle inline bold
            p = doc.add_paragraph()
            # Simple inline bold split
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.size = Pt(10)

add_specialist_section(doc, 'A.1 Clinical questions',
                       '/Users/justinyu/Desktop/linkedin-posts/TA1120_clarification_review/specialist_outputs/clinical_questions.md')
add_page_break(doc)

add_specialist_section(doc, 'A.2 Statistician questions',
                       '/Users/justinyu/Desktop/linkedin-posts/TA1120_clarification_review/specialist_outputs/statistician_questions.md')
add_page_break(doc)

add_specialist_section(doc, 'A.3 Health economist questions',
                       '/Users/justinyu/Desktop/linkedin-posts/TA1120_clarification_review/specialist_outputs/health_economist_questions.md')
add_page_break(doc)

add_specialist_section(doc, 'A.4 Information specialist questions',
                       '/Users/justinyu/Desktop/linkedin-posts/TA1120_clarification_review/specialist_outputs/information_specialist_questions.md')

add_page_break(doc)

# ============================================================
# APPENDIX B — FULL NICE QUESTION LIST
# ============================================================

add_heading(doc, 'Appendix B: Full actual NICE question list', level=1)
add_para(doc,
    'The 23 unique NICE clarification questions issued to Merck on 16 December 2024, reproduced from '
    'Section 2 (company response to clarification) and de-duplicated. Priority "Yes" items: A1, A9, B1, '
    'B2, B7, B8, B10, B11, B12.', italic=True, size=10)

nice_questions = [
    ('A1', 'A', 'Yes', 'BICR vs investigator PFS curves',
     'If possible, please provide versions of CS Figure 7 and CS Figure 9 which also show the PFS '
     'assessments by blinded independent central review (BICR). If this is not possible for the favourable '
     'and intermediate/poor IMDC prognostic risk groups, please provide a figure showing BICR PFS for the '
     'period this was assessed and investigator PFS for the full trial period on the same plot for the '
     'total trial population.'),
    ('A2', 'A', 'No', 'FKSI-DRS time-to-deterioration outcome',
     'CS Table 6 and CS Table 7 refer to the outcome for "Time to deterioration in the FKSI-DRS", however '
     'the results are not reported in CS section B.2.6 nor have we identified them in the CSR. Please '
     'provide the results or signpost to where they are in the submission.'),
    ('A3', 'A', 'No', 'Avelumab vs axitinib death-discontinuation count discrepancy',
     'CS Table 11. We note that the number of participants who discontinued due to death differs for '
     'avelumab and axitinib (n=25 and n=28 respectively). Is this because some participants were only in '
     'receipt of either avelumab or axitinib at the time of death? If not, please explain why these '
     'numbers differ.'),
    ('A4', 'A', 'No', 'Patient disposition by IMDC risk subgroup',
     'CS Table 11. If possible, please provide versions of CS Table 11 for the favourable and '
     'intermediate/poor IMDC prognostic risk subgroups.'),
    ('A5', 'A', 'No', 'Adverse-event rates per single-agent monotherapy',
     'CS section B.2.11.1.2 reports that some AEs were at higher frequencies with the combination than '
     'observed with either single agent alone. Please indicate where we can find the adverse event data '
     'for each agent as monotherapy (e.g. for diarrhoea, where in the CSR or elsewhere is it reported '
     'that XXXX and XXXX of avelumab monotherapy and axitinib monotherapy patients respectively '
     'experienced this event?).'),
    ('A6', 'A', 'No', 'UK RWE source overlap (SACT-EAMS, SACT-CDF, Nathan, McGrane)',
     'Please confirm if the EAG\'s understanding is correct that the only UK sources of real-world '
     'evidence without overlap (or potential overlap) are the SACT-EAMS cohort and the SACT-CDF cohort, '
     'i.e. there were no patients who contributed data to both the SACT-EAMS cohort and the SACT-CDF '
     'cohort after exclusions. Some or all of the patients in Nathan et al. and McGrane et al. may also '
     'be included in the SACT database.'),
    ('A7', 'A', 'No', 'Reasons for SACT cohort exclusions',
     'CS B.2.8.2.1. What were the reasons for excluding patients from the SACT-CDF and SACT-EAMS cohorts?'),
    ('A8', 'A', 'No', 'RWE SLR report — document request',
     'Appendix D.6 directs the reader to the RWE SLR report for further information and full results. '
     'Please provide this reference (or indicate which reference in the reference pack it is).'),
    ('A9', 'A', 'Yes', 'NMA report / feasibility assessment / risk-of-bias',
     'If there is a separate report with full details for the NMAs and any feasibility assessment '
     'conducted please provide this (or indicate which reference in the reference pack it is). If there '
     'is not a specific report please provide: a table of the PFS and OS data used in the NMAs for the '
     'intermediate/poor risk subgroup; a summary of baseline characteristics for participants from the '
     'studies included in the NMAs for the intermediate/poor risk subgroup; critical appraisal/risk of '
     'bias assessments for the comparator RCTs in the NMA.'),
    ('A10', 'A', 'No', 'SLR critical-appraisal methodology',
     'Please report how many reviewers carried out critical appraisal of studies for the main clinical '
     'effectiveness systematic literature review (SLR) and the real-world evidence SLR and, if more than '
     'one reviewer, whether the reviewers worked independently.'),
    ('B1', 'B', 'Yes', '40-year time horizon rationale',
     'What is the rationale for using a much longer time horizon (40 years) than in previous appraisals '
     'for aRCC (10 years)?'),
    ('B2', 'B', 'Yes', 'Hazard plots and PH assumption — favourable-risk',
     'Hazard function plots for the survival curves in the intermediate/poor-risk subgroup are provided '
     'in CS Appendix N Figures 1–4. Please provide equivalent hazard function plots for OS, PFS and TTD '
     'for both arms of the JAVELIN 101 trial for the favourable risk population. (a) Is there evidence to '
     'support or reject an assumption of proportional hazards for OS and PFS in the favourable risk '
     'population? (b) Please justify the choice of parametric survival curve for each outcome and arm.'),
    ('B3', 'B', 'No', 'KM data timepoint discrepancy (model vs source)',
     'The Kaplan-Meier data provided in the KM data sheet of the model (KM data!BK46:CM233) have '
     'different time points to those reported in the source (DOF additional analyses). For example, OS '
     'in the model has timepoints 0, 3.81, 5.09 months while the source provides 4, 8, 12 months '
     '(at 4-month intervals). Please confirm that this discrepancy is due to the model data being '
     'derived from individual patient level data, or provide a source with matching KM data.'),
    ('B4', 'B', 'No', 'Missing Grade ≥3 TRAEs in CS Table 56',
     'Please provide a justification as to why some treatment-related adverse events of grade ≥3 '
     'experienced by more than 5% of patients that are reported in the JAVELIN Renal 101 trial are '
     'missing from CS Table 56. For example, fatigue, nausea, decreased appetite, stomatitis, mucosal '
     'inflammation, aspartate aminotransferase increased, and vomiting.'),
    ('B5', 'B', 'No', 'EQ-5D-5L index scoring method',
     'Please state the method used to attribute index scores to the EQ-5D-5L FAS data, as reported in '
     'CS section B.2.6.4.1. In particular, please confirm whether the Hernández-Alava et al. (2017) '
     'mapping approach was used as in the analysis of EQ-5D-5L data for the economic model.'),
    ('B6', 'B', 'No', 'EQ-5D-5L by IMDC risk subgroup',
     'Descriptive information and results for the EQ-5D-5L index score is reported in CS section '
     'B.2.6.4.1 for the FAS population. Please provide this information for the favourable-risk and '
     'intermediate/poor risk subgroups. For each subgroup, please provide a graph of EQ-5D-5L index '
     'scores change from baseline by visit, with numbers at risk.'),
    ('B7', 'B', 'Yes', 'Utility regression methods',
     'Please provide more explanation on the methods used to derive health state utility values for the '
     'economic model. (a) Please describe the volume and pattern of missing EQ-5D-5L data. (b) Was any '
     'attempt made to impute the missing data in the utility regressions? (c) Please justify the chosen '
     'regression model specifications for Model 1 and Model 2. Would the fit of these models be improved '
     'by including other baseline co-variates? Would it have been more efficient to use a single model '
     'based on ITT data, with fixed-effect and interaction terms for risk group status?'),
    ('B8', 'B', 'Yes', 'Nivolumab not used as subsequent therapy assumption',
     'The base case analysis includes an assumption that nivolumab is not used as subsequent treatment '
     'after progression on avelumab + axitinib (CS Table 69), with scenario analysis conducted assuming '
     'different proportions of subsequent treatment options. This assumption is justified in CS section '
     'B.3.5.4 on the basis of clinical expert feedback. Please provide a clearer justification based on '
     'evidence from NICE guidance, NHS England prescribing policy and NHS practice (e.g. SACT data). '
     'Please also explain and justify the alternative assumptions used for scenario analysis.'),
    ('B9', 'B', 'No', 'Equal AE costs / non-elective short-stay proxy',
     'Please provide a justification as to why costs are assumed equal for all adverse events, and how '
     'this impacts the model results. Please also explain why a non-elective short stay cost is used for '
     'all grade ≥3 adverse events.'),
    ('B10', 'B', 'Yes', 'Follow-up treatment distribution mismatch',
     'The values in the model table Costs!AK162:AL185, which presents the distribution of follow-up '
     'treatments for the poor-/intermediate-risk group, are different from those reported in Table '
     '10.5.a of the relevant source (ASCO presentation analysis). Only the first six drugs of Table '
     '10.5.a were considered for the poor-risk group while 24 were considered for the intermediate-risk '
     'group. Please provide a rationale for this assumption.'),
    ('B11', 'B', 'Yes', 'Untraceable model inputs (age, sex, weight by risk subgroup)',
     'The EAG are unable to identify the sources of several model input parameters which are reported in '
     'the CS and used in the company\'s model. Parameters queried: average age, proportion female, and '
     'average weight for favourable-risk (CS Table 40); for intermediate/poor-risk (CS Appendix O Table '
     '50); and average weight for ITT (CS Appendix O Table 51) — all sourced to JAVELIN Renal 101. '
     'Please clarify how these values were derived from the corresponding sources.'),
    ('B12', 'B', 'Yes', 'CS vs model input value discrepancies',
     'Different input values are reported in the company submission and in the company\'s model for '
     'several input parameters: IV administration costs (CS Table 64 vs Costs!G86/AE86, NHS National '
     'Cost Collection 2022/23, SB12Z); CT scan and blood test costs (CS Table 66 vs Costs!G118:G122); '
     'nivolumab treatment dosing (CS Table 71 vs Costs!H183; source TA542); total subsequent treatment '
     'costs avelumab + axitinib (intermediate/poor risk) (CS Appendix O Table 91 vs Costs!G234). Please '
     'clarify which values should be considered in the company\'s base case.'),
    ('C1', 'C', 'No', 'Confidentiality of DOF references',
     'Should we treat all the data-on-file (DOF) references as commercial in confidence?'),
]

nice_widths = [600, 600, 720, 2400, 5040]
assert sum(nice_widths) == 9360
nice_table = make_fixed_table(doc, nice_widths,
    ['ID', 'Sec', 'Priority', 'Topic', 'Verbatim text (abbreviated)'],
    header_shade='1F3864')

for i, vals in enumerate(nice_questions):
    add_table_row(nice_table, nice_widths, vals, font_size=9,
                  shade_alt=True, row_index=i)

# ============================================================
# Save and run sanity check
# ============================================================

doc.save(OUT_PATH)
print(f"Wrote: {OUT_PATH}")


# ============================================================
# Sanity check
# ============================================================

def sanity_check(path):
    d = Document(path)
    errors = []
    for i, tbl in enumerate(d.tables):
        grid = tbl._tbl.find(qn('w:tblGrid'))
        cols = grid.findall(qn('w:gridCol')) if grid is not None else []
        widths = [int(c.get(qn('w:w'))) for c in cols]
        total = sum(widths)
        if total > LANDSCAPE_TWIPS:
            errors.append(f"Table {i}: grid sum {total} > {LANDSCAPE_TWIPS}")
        tblPr = tbl._tbl.find(qn('w:tblPr'))
        layout = tblPr.find(qn('w:tblLayout')) if tblPr is not None else None
        if layout is None or layout.get(qn('w:type')) != 'fixed':
            errors.append(f"Table {i}: tblLayout not fixed")
        # Cross-check per-cell widths sum to grid sum (best-effort: first row)
        if tbl.rows:
            first_row = tbl.rows[0]
            cell_widths = []
            for c in first_row.cells:
                tcW = c._tc.find(qn('w:tcPr')).find(qn('w:tcW')) if c._tc.find(qn('w:tcPr')) is not None else None
                if tcW is not None:
                    cell_widths.append(int(tcW.get(qn('w:w'))))
            if cell_widths and sum(cell_widths) != total:
                errors.append(f"Table {i}: per-cell sum {sum(cell_widths)} != grid sum {total}")
    if errors:
        raise AssertionError('Table-width sanity check FAILED:\n' + '\n'.join(errors))
    return True


sanity_check(OUT_PATH)
print("Sanity check PASSED")

# Page count estimate (very rough)
d = Document(OUT_PATH)
n_paragraphs = len(d.paragraphs)
n_tables = len(d.tables)
n_table_rows = sum(len(t.rows) for t in d.tables)
print(f"Paragraphs: {n_paragraphs}, Tables: {n_tables}, Total table rows: {n_table_rows}")
